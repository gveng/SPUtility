from __future__ import annotations

from copy import deepcopy
from importlib import resources
from io import BytesIO
import json
import shutil
from datetime import datetime
from pathlib import Path
from uuid import uuid4
from weakref import ref

import numpy as np

from PySide6.QtCore import QBuffer, QByteArray, QIODevice, QSettings, Qt, QUrl
from PySide6.QtGui import QCloseEvent, QColor, QPainter, QPen, QPixmap
from PySide6.QtWidgets import (
    QApplication,
    QDialog,
    QFileDialog,
    QLabel,
    QMainWindow,
    QMenu,
    QMessageBox,
    QScrollArea,
    QStyle,
    QTextBrowser,
    QTreeWidget,
    QTreeWidgetItem,
    QVBoxLayout,
    QWidget,
)

import sparams_utility as pkg
from sparams_utility.circuit_solver import ChannelSimResult, TransientSimResult, TransientTrace
from sparams_utility.models.circuit import DriverSpec, TransientSourceSpec
from sparams_utility.models.state import AppState, LoadedTouchstone
from sparams_utility.touchstone_parser import parse_touchstone_string
from sparams_utility.ui.child_windows import ChildWindowManager, default_child_window_size
from sparams_utility.ui.circuit_window import CircuitWindow
from sparams_utility.ui.eye_diagram_window import EyeDiagramWindow
from sparams_utility.ui.plot_window import PlotWindow
from sparams_utility.ui.tdr_window import TdrWindow
from sparams_utility.ui.table_models import MagnitudeTableModel, RawDataTableModel
from sparams_utility.ui.table_window import TableWindow
from sparams_utility.ui.transient_window import TransientResultWindow


_RECENT_ENTRIES_MAX = 12


class MainWindow(QMainWindow):
    def __init__(self, state: AppState) -> None:
        super().__init__()
        self.setWindowTitle(f"{pkg.__app_name__}  {pkg.__version__}")
        self.resize(360, 800)
        app = QApplication.instance()
        if app is not None:
            self.setWindowIcon(app.windowIcon())

        self._state = state
        self._plot_counter: int = 0
        self._tdr_counter: int = 0
        self._circuit_counter: int = 0
        self._project_dirty: bool = False
        self._project_path: str | None = None
        self._recent_projects: list[str] = []
        self._recent_sparams: list[str] = []
        self._project_data_dir: Path | None = None
        self._window_registry: dict[str, dict[str, dict]] = {
            "sp": {},
            "tdr": {},
            "circuit": {},
        }
        self._open_windows: dict[tuple[str, str], object] = {}
        self._open_window_keys: dict[int, tuple[str, str]] = {}
        self._load_recent_entries()

        # Child window manager — top-level category windows replace the MDI area.
        self._windows = ChildWindowManager(
            self,
            icon_for_category=self._icon_for_category,
            icon_for_widget=self._icon_for_widget,
        )
        self._windows.widget_activated.connect(self._on_active_widget_changed)
        self._windows.widget_closed.connect(self._on_managed_widget_closed)

        self._project_tree = QTreeWidget()
        self._project_tree.setHeaderHidden(True)
        self._project_tree.setMinimumWidth(220)
        self._project_tree.itemClicked.connect(self._on_project_tree_item_clicked)
        self._project_tree.setContextMenuPolicy(Qt.CustomContextMenu)
        self._project_tree.customContextMenuRequested.connect(self._on_project_tree_context_menu)

        tree_panel = QWidget()
        tree_layout = QVBoxLayout(tree_panel)
        tree_layout.setContentsMargins(6, 6, 6, 6)
        tree_layout.setSpacing(4)
        tree_layout.addWidget(QLabel("Project"))
        tree_layout.addWidget(self._project_tree)

        self.setCentralWidget(tree_panel)

        # ── File ──────────────────────────────────────────────────────────
        file_menu = self.menuBar().addMenu("File")
        file_menu.addAction("Open Project", self._open_project)
        file_menu.addAction("Save Project", self._save_project)
        file_menu.addAction("Save Project As", self._save_project_as)
        file_menu.addAction("Close Project", self._close_project)
        file_menu.addAction("Export Project", self._export_project)
        self._recent_projects_menu = file_menu.addMenu("Recent Projects")
        file_menu.addSeparator()
        file_menu.addAction("Open File", self._open_files)
        self._recent_sparams_menu = file_menu.addMenu("Recent S-Parameters")
        self._rebuild_recent_file_menus()
        file_menu.addSeparator()
        file_menu.addAction("Exit", self.close)

        # ── Tables ────────────────────────────────────────────────────────
        self._tables_menu = self.menuBar().addMenu("Tables")

        # ── Charts ────────────────────────────────────────────────────────
        charts_menu = self.menuBar().addMenu("Charts")
        charts_menu.addAction("Open New SP Plot Window", self._open_plot_window)
        charts_menu.addAction("Open new TDR Window", self._open_tdr_window)

        report_menu = self.menuBar().addMenu("Report")
        report_menu.addAction("Create SP Plots Reports", self._create_sp_plots_report)
        report_menu.addAction("Create TDR plots reports", self._create_tdr_plots_report)
        report_menu.addAction("Create EYE diagrams report", self._create_eye_diagrams_report)

        circuit_menu = self.menuBar().addMenu("Circuit")
        circuit_menu.addAction("Open circuit composer", self._open_circuit_window)

        self._active_circuit_menu = self.menuBar().addMenu("Circuit Window")
        self._active_circuit_menu.addAction("Duplicate This circuit...", self._duplicate_active_circuit)
        self._active_circuit_menu.menuAction().setVisible(False)

        # ── View ──────────────────────────────────────────────────────────
        view_menu = self.menuBar().addMenu("View")
        view_menu.addAction("Minimize all", self._minimize_all)
        view_menu.addAction("Restore all", self._restore_all)
        view_menu.addAction("Resize all graph with last graph window", self._resize_all_graph_windows)
        view_menu.addSeparator()
        view_menu.addAction("Close all", self._windows.close_all)

        # ── Window ────────────────────────────────────────────────────────
        self._window_menu = self.menuBar().addMenu("Window")
        self._window_menu.aboutToShow.connect(self._populate_window_menu)

        # ── Help ──────────────────────────────────────────────────────────
        help_menu = self.menuBar().addMenu("Help")
        help_menu.addAction("User Guide", self._show_help)
        help_menu.addAction("About", self._show_about)

        self._state.files_changed.connect(self._rebuild_tables_menu)
        self._state.files_changed.connect(self._refresh_project_tree)
        self._rebuild_tables_menu()
        self._on_active_widget_changed(None)
        self._refresh_project_tree()

    def _project_tree_project_label(self) -> str:
        if self._project_path:
            return self._tree_display_name(Path(self._project_path).name)
        return "Untitled Project"

    def _tree_display_name(self, name: str) -> str:
        raw_name = str(name).strip()
        if not raw_name:
            return raw_name
        suffix = Path(raw_name).suffix
        if suffix:
            return Path(raw_name).stem or raw_name
        return raw_name

    def _icon_for_kind(self, kind: str):
        style = self.style()

        def _load_custom_icon(resource_name: str):
            try:
                import importlib.resources as pkg_resources
                from PySide6.QtGui import QIcon
                import sparams_utility.resources

                resource = pkg_resources.files(sparams_utility.resources).joinpath(resource_name)
                with pkg_resources.as_file(resource) as path:
                    icon = QIcon(str(path))
                    if not icon.isNull():
                        return icon
            except Exception:
                pass
            return None

        if kind == "tdr":
            icon = _load_custom_icon("pulse_icon.svg")
            if icon is not None:
                return icon
            return style.standardIcon(QStyle.SP_DriveNetIcon)
        if kind == "circuit":
            icon = _load_custom_icon("circuit_icon.svg")
            if icon is not None:
                return icon
            # Fallback to default directory icon
            return style.standardIcon(QStyle.SP_DirIcon)
        if kind == "eye-file":
            icon = _load_custom_icon("eye_icon.svg")
            if icon is not None:
                return icon
            # Fallback to default detailed view icon
            return style.standardIcon(QStyle.SP_FileDialogDetailedView)
        if kind == "transient-plot":
            icon = _load_custom_icon("sine_wave.svg")
            if icon is not None:
                return icon
            return style.standardIcon(QStyle.SP_FileDialogDetailedView)
        if kind == "transient-file":
            icon = _load_custom_icon("sine_wave.svg")
            if icon is not None:
                return icon
            return style.standardIcon(QStyle.SP_FileDialogDetailedView)
        if kind in {"sp", "sparam-file", "touchstone-file"}:
            icon = _load_custom_icon("cosine_wave.svg")
            if icon is not None:
                return icon
        if kind == "sp":
            return style.standardIcon(QStyle.SP_ComputerIcon)
        if kind == "sparam-file":
            # Fallback to default file link icon
            return style.standardIcon(QStyle.SP_FileLinkIcon)
        if kind == "touchstone-file":
            return style.standardIcon(QStyle.SP_FileIcon)
        return style.standardIcon(QStyle.SP_FileIcon)

    # ── Window taskbar icons ──────────────────────────────────────────────
    def _icon_for_category(self, category: str):
        """Icon for a CategoryWindow (taskbar grouping by element type)."""
        mapping = {
            "circuits": "circuit",
            "plots": "sp",
            "tdr": "tdr",
            "eye": "eye-file",
            "transient": "transient-plot",
            "tables": "sparam-file",
        }
        kind = mapping.get(category)
        if kind is not None:
            return self._icon_for_kind(kind)
        # Fallback: app icon
        app = QApplication.instance()
        if app is not None:
            return app.windowIcon()
        return self.style().standardIcon(QStyle.SP_FileIcon)

    def _icon_for_widget(self, widget):
        """Per-window taskbar icon, matching the project-tree element type."""
        if isinstance(widget, CircuitWindow):
            return self._icon_for_kind("circuit")
        if isinstance(widget, PlotWindow):
            return self._icon_for_kind("sp")
        if isinstance(widget, TdrWindow):
            return self._icon_for_kind("tdr")
        if isinstance(widget, EyeDiagramWindow):
            return self._icon_for_kind("eye-file")
        if isinstance(widget, TransientResultWindow):
            return self._icon_for_kind("transient-plot")
        if isinstance(widget, TableWindow):
            return self._icon_for_kind("sparam-file")
        app = QApplication.instance()
        if app is not None:
            return app.windowIcon()
        return self.style().standardIcon(QStyle.SP_FileIcon)

    def _new_entry_id(self, kind: str) -> str:
        return f"{kind}:{uuid4().hex[:10]}"

    def _register_window_entry(
        self,
        kind: str,
        widget: object,
        *,
        state: dict | None = None,
        parent_circuit_entry_id: str | None = None,
    ) -> str:
        entry_id = self._new_entry_id(kind)
        window_number = int(getattr(widget, "window_number", 0))
        self._window_registry[kind][entry_id] = {
            "id": entry_id,
            "kind": kind,
            "title": str(widget.windowTitle()),
            "state": state if isinstance(state, dict) else self._snapshot_widget_state(widget),
            "window_number": window_number,
            "is_open": False,
            "eye_file": None,
            "sparam_file": None,
            "sparam_plot_file": None,
            "transient_file": None,
            "output_kind": None,
            "parent_circuit_entry_id": parent_circuit_entry_id,
        }
        return entry_id

    def _snapshot_widget_state(self, widget: object) -> dict:
        if isinstance(widget, PlotWindow):
            return widget.export_project_state()
        if isinstance(widget, TdrWindow):
            return widget.export_project_state()
        if isinstance(widget, CircuitWindow):
            return widget.export_project_state()
        return {}

    def _bind_open_window(self, kind: str, entry_id: str, widget: object) -> None:
        entry = self._window_registry[kind].get(entry_id)
        if entry is None:
            return
        entry["is_open"] = True
        entry["title"] = str(widget.windowTitle())
        entry["window_number"] = int(getattr(widget, "window_number", entry.get("window_number", 0)))
        entry["state"] = self._snapshot_widget_state(widget)
        entry["window_size"] = [widget.width(), widget.height()]
        self._open_windows[(kind, entry_id)] = widget
        self._open_window_keys[id(widget)] = (kind, entry_id)

        if hasattr(widget, "project_modified"):
            widget.project_modified.connect(
                lambda k=kind, eid=entry_id, w=widget: self._sync_window_entry(k, eid, w)
            )
        if isinstance(widget, PlotWindow):
            widget.unload_file_requested.connect(self._unload_file)
        if isinstance(widget, CircuitWindow):
            widget.eye_result_generated.connect(
                lambda payload, eid=entry_id: self._on_circuit_eye_result_generated(eid, payload)
            )
            widget.sparameter_result_generated.connect(
                lambda payload, eid=entry_id: self._on_circuit_sparameter_result_generated(eid, payload)
            )
            widget.transient_result_generated.connect(
                lambda payload, eid=entry_id: self._on_circuit_transient_result_generated(eid, payload)
            )
            widget.transient_windows_changed.connect(self._refresh_project_tree)
        owner_ref = ref(self)
        widget.destroyed.connect(
            lambda *_, owner_ref=owner_ref, kind=kind, entry_id=entry_id: MainWindow._on_window_closed_safe(
                owner_ref,
                kind,
                entry_id,
            )
        )

    @staticmethod
    def _on_window_closed_safe(owner_ref, kind: str, entry_id: str) -> None:
        owner = owner_ref()
        if owner is None:
            return
        try:
            owner._on_window_closed(kind, entry_id)
        except RuntimeError:
            return

    def _ensure_data_dir_for_runtime_exports(self) -> Path | None:
        if self._project_data_dir is not None:
            try:
                self._project_data_dir.mkdir(parents=True, exist_ok=True)
            except OSError:
                return None
            return self._project_data_dir

        if self._project_path:
            self._project_data_dir = Path(self._project_path).with_name(Path(self._project_path).stem + "_Data")
        else:
            fallback = f"{pkg.__app_name__.replace(' ', '_')}_Data"
            self._project_data_dir = Path.cwd() / fallback

        try:
            self._project_data_dir.mkdir(parents=True, exist_ok=True)
        except OSError:
            return None
        return self._project_data_dir

    def _build_eye_binary_payload(
        self,
        result: ChannelSimResult,
        *,
        eye_span_ui: int,
        render_mode: str,
        quality_preset: str,
        stat_enabled: bool,
        noise_rms_mv: float,
        jitter_rms_ps: float,
    ) -> bytes:
        payload = BytesIO()
        np.savez_compressed(
            payload,
            time_s=np.asarray(result.time_s, dtype=float),
            waveform_v=np.asarray(result.waveform_v, dtype=float),
            ui_s=np.asarray([float(result.ui_s)], dtype=float),
            is_differential=np.asarray([1 if result.is_differential else 0], dtype=np.int8),
            driver_spec_json=np.asarray([json.dumps(result.driver_spec.to_dict())]),
            eye_span_ui=np.asarray([int(eye_span_ui)], dtype=np.int32),
            render_mode=np.asarray([str(render_mode)]),
            quality_preset=np.asarray([str(quality_preset)]),
            stat_enabled=np.asarray([1 if stat_enabled else 0], dtype=np.int8),
            noise_rms_mv=np.asarray([float(noise_rms_mv)], dtype=float),
            jitter_rms_ps=np.asarray([float(jitter_rms_ps)], dtype=float),
        )
        return payload.getvalue()

    def _on_circuit_eye_result_generated(self, entry_id: str, payload: object) -> None:
        entry = self._window_registry.get("circuit", {}).get(entry_id)
        if entry is None or not isinstance(payload, dict):
            return

        result = payload.get("result")
        if not isinstance(result, ChannelSimResult):
            return

        data_dir = self._ensure_data_dir_for_runtime_exports()
        if data_dir is None:
            return

        circuit_name = str(payload.get("circuit_name") or entry.get("title") or f"Circuit_{entry_id}")
        eye_file_name = self._sanitize_file_stem(circuit_name, f"Circuit_{entry_id}") + "_Eye.eye"
        eye_path = data_dir / eye_file_name

        try:
            binary_payload = self._build_eye_binary_payload(
                result,
                eye_span_ui=int(payload.get("eye_span_ui", 2)),
                render_mode=str(payload.get("render_mode", "Density")),
                quality_preset=str(payload.get("quality_preset", "Balanced")),
                stat_enabled=bool(payload.get("stat_enabled", False)),
                noise_rms_mv=float(payload.get("noise_rms_mv", 0.0)),
                jitter_rms_ps=float(payload.get("jitter_rms_ps", 0.0)),
            )
            eye_path.write_bytes(binary_payload)
        except Exception:
            return

        old_eye_file = str(entry.get("eye_file") or "").strip()
        old_sparam_file = str(entry.get("sparam_file") or "").strip()
        old_sparam_plot_file = str(entry.get("sparam_plot_file") or "").strip()
        entry["eye_file"] = eye_file_name
        entry["sparam_file"] = None
        entry["sparam_plot_file"] = None
        entry["output_kind"] = "eye"
        if old_eye_file and old_eye_file != eye_file_name:
            old_path = data_dir / old_eye_file
            if old_path.exists():
                try:
                    old_path.unlink()
                except OSError:
                    pass
        for stale_name in (old_sparam_file, old_sparam_plot_file):
            if stale_name:
                stale_path = data_dir / stale_name
                if stale_path.exists():
                    try:
                        stale_path.unlink()
                    except OSError:
                        pass
        self._mark_project_dirty()
        self._refresh_project_tree()

    def _on_circuit_sparameter_result_generated(self, entry_id: str, payload: object) -> None:
        entry = self._window_registry.get("circuit", {}).get(entry_id)
        if entry is None or not isinstance(payload, dict):
            return

        touchstone_text = str(payload.get("touchstone_text") or "")
        nports = int(payload.get("nports") or 0)
        circuit_name = str(payload.get("circuit_name") or entry.get("title") or f"Circuit_{entry_id}")
        if not touchstone_text or nports <= 0:
            payload["error"] = "Generated Touchstone payload is incomplete."
            return

        data_dir = self._ensure_data_dir_for_runtime_exports()
        if data_dir is None:
            payload["error"] = "Project data folder is not available."
            return

        sparam_file_name = self._normalize_generated_touchstone_file_name(
            str(payload.get("file_name") or ""),
            nports,
            self._sanitize_file_stem(circuit_name, f"Circuit_{entry_id}"),
        )
        sparam_path = data_dir / sparam_file_name

        try:
            sparam_path.write_text(touchstone_text, encoding="utf-8")
        except Exception as exc:
            payload["error"] = f"Could not save Touchstone file: {exc}"
            return

        loaded, error = self._state.load_or_reload_file(str(sparam_path))
        if loaded is None:
            payload["error"] = error or "Could not load generated Touchstone file."
            return

        self._remove_generated_plot_entries_for_circuit(entry_id)

        old_eye_file = str(entry.get("eye_file") or "").strip()
        old_sparam_file = str(entry.get("sparam_file") or "").strip()
        old_sparam_plot_file = str(entry.get("sparam_plot_file") or "").strip()
        entry["eye_file"] = None
        entry["sparam_file"] = sparam_file_name
        entry["sparam_plot_file"] = None
        entry["output_kind"] = "sparam"

        if old_eye_file:
            old_eye_path = data_dir / old_eye_file
            if old_eye_path.exists():
                try:
                    old_eye_path.unlink()
                except OSError:
                    pass
        if old_sparam_file and old_sparam_file != sparam_file_name:
            self._state.unload_file(str((data_dir / old_sparam_file).resolve()))

        for old_name, new_name in ((old_sparam_file, sparam_file_name), (old_sparam_plot_file, "")):
            if old_name and old_name != new_name:
                old_path = data_dir / old_name
                if old_path.exists():
                    try:
                        old_path.unlink()
                    except OSError:
                        pass

        plot_entry_id = self._open_generated_sparameter_plot_window(
            loaded,
            parent_circuit_entry_id=entry_id,
        )
        payload["saved_file_name"] = sparam_file_name
        payload["saved_path"] = str(sparam_path)
        payload["plot_entry_id"] = plot_entry_id

        self._mark_project_dirty()
        self._refresh_project_tree()

    def _normalize_generated_touchstone_file_name(
        self,
        requested_name: str,
        nports: int,
        fallback_stem: str,
    ) -> str:
        raw_name = Path(str(requested_name).strip()).name
        stem = Path(raw_name).stem if raw_name else fallback_stem
        safe_stem = self._sanitize_file_stem(stem, fallback_stem)
        return f"{safe_stem}.s{nports}p"

    def _build_generated_plot_state(self, loaded: LoadedTouchstone) -> dict:
        plot_name = self._tree_display_name(loaded.display_name)
        excluded_files = [
            str(item.path)
            for item in self._state.get_loaded_files()
            if item.file_id != loaded.file_id
        ]
        return {
            "window_title": plot_name,
            "plot_name": plot_name,
            "plot_settings": {},
            "legend_position": [10.0, 10.0],
            "files": [
                {
                    "file_path": str(loaded.path),
                    "file_name": loaded.display_name,
                    "legend_label": loaded.display_name,
                    "selected_parameters": [],
                }
            ],
            "excluded_files": excluded_files,
        }

    def _open_generated_sparameter_plot_window(
        self,
        loaded: LoadedTouchstone,
        *,
        parent_circuit_entry_id: str,
    ) -> str:
        self._plot_counter += 1
        plot_win = PlotWindow(self._state, window_number=self._plot_counter)
        plot_win.apply_project_state(self._build_generated_plot_state(loaded))
        plot_win.project_modified.connect(self._mark_project_dirty)
        plot_win.resize(*default_child_window_size("plots"))
        entry_id = self._register_window_entry(
            "sp",
            plot_win,
            parent_circuit_entry_id=parent_circuit_entry_id,
        )
        self._bind_open_window("sp", entry_id, plot_win)
        self._windows.present(plot_win)
        self._windows.bring_to_front(plot_win)
        return entry_id

    def _remove_generated_plot_entries_for_circuit(self, circuit_entry_id: str) -> None:
        related_plot_ids = [
            plot_entry_id
            for plot_entry_id, plot_entry in self._window_registry["sp"].items()
            if str(plot_entry.get("parent_circuit_entry_id") or "") == circuit_entry_id
        ]
        for plot_entry_id in related_plot_ids:
            self._delete_tree_window_entry("sp", plot_entry_id)

    def _sync_window_entry(self, kind: str, entry_id: str, widget: object) -> None:
        entry = self._window_registry[kind].get(entry_id)
        if entry is None:
            return
        entry["title"] = str(widget.windowTitle())
        entry["state"] = self._snapshot_widget_state(widget)
        entry["is_open"] = True
        if hasattr(widget, "size"):
            size = widget.size()
            entry["window_size"] = [size.width(), size.height()]
        self._refresh_project_tree()

    def _on_window_closed(self, kind: str, entry_id: str) -> None:
        widget = self._open_windows.pop((kind, entry_id), None)
        if widget is not None:
            try:
                self._open_window_keys.pop(id(widget), None)
            except RuntimeError:
                pass
        entry = self._window_registry[kind].get(entry_id)
        if entry is not None:
            entry["is_open"] = False
        self._refresh_project_tree()

    def _on_managed_widget_closed(self, widget) -> None:
        try:
            key = self._open_window_keys.get(id(widget))
        except RuntimeError:
            return
        if key is None:
            return
        kind, entry_id = key
        self._on_window_closed(kind, entry_id)

    def _add_tree_window_item(self, parent: QTreeWidgetItem, kind: str, entry_id: str, title: str) -> QTreeWidgetItem:
        entry = self._window_registry[kind].get(entry_id, {})
        label = title.strip()
        if kind == "circuit" and " - " in label:
            label = label.split(" - ", 1)[1].strip() or label
        if kind == "sp" and label.startswith("S-Parameter Plots #"):
            suffix = label.split("#", 1)[1].strip() if "#" in label else ""
            label = f"Plot #{suffix}" if suffix else "Plot"
        if kind == "tdr" and label.startswith("TDR Plots #"):
            suffix = label.split("#", 1)[1].strip() if "#" in label else ""
            label = f"TDR #{suffix}" if suffix else "TDR"

        item = QTreeWidgetItem([label])
        item.setIcon(0, self._icon_for_kind(kind))
        item.setData(0, Qt.UserRole, {"action": "window", "kind": kind, "id": entry_id})
        parent.addChild(item)

        if kind == "circuit":
            eye_file = str(entry.get("eye_file") or "").strip()
            if eye_file:
                eye_item = QTreeWidgetItem([self._tree_display_name(eye_file)])
                eye_item.setIcon(0, self._icon_for_kind("eye-file"))
                eye_item.setData(
                    0,
                    Qt.UserRole,
                    {
                        "action": "output-file",
                        "kind": kind,
                        "id": entry_id,
                        "file": eye_file,
                        "output": "eye",
                    },
                )
                item.addChild(eye_item)
                item.setExpanded(True)

            sparam_plot_file = str(entry.get("sparam_plot_file") or "").strip()
            if sparam_plot_file:
                s_item = QTreeWidgetItem([self._tree_display_name(sparam_plot_file)])
                s_item.setIcon(0, self._icon_for_kind("sparam-file"))
                s_item.setData(
                    0,
                    Qt.UserRole,
                    {
                        "action": "output-file",
                        "kind": kind,
                        "id": entry_id,
                        "file": sparam_plot_file,
                        "output": "sparam-plot",
                    },
                )
                item.addChild(s_item)
                item.setExpanded(True)

            circuit_widget = self._open_windows.get(("circuit", entry_id))

            transient_file = str(entry.get("transient_file") or "").strip()
            if transient_file:
                t_item = QTreeWidgetItem([self._tree_display_name(transient_file)])
                t_item.setIcon(0, self._icon_for_kind("transient-file"))
                t_item.setData(
                    0,
                    Qt.UserRole,
                    {
                        "action": "output-file",
                        "kind": kind,
                        "id": entry_id,
                        "file": transient_file,
                        "output": "transient",
                    },
                )
                item.addChild(t_item)
                item.setExpanded(True)

            return item

    def _refresh_project_tree(self) -> None:
        self._project_tree.clear()

        project_item = QTreeWidgetItem([self._project_tree_project_label()])
        self._project_tree.addTopLevelItem(project_item)

        flat_entries: list[tuple[str, str, str]] = []
        nested_plot_entries: dict[str, list[tuple[str, str, str]]] = {}
        for kind in ("circuit", "sp", "tdr"):
            for entry in self._window_registry[kind].values():
                entry_id = str(entry.get("id", ""))
                title = str(entry.get("title", "Window"))
                if kind == "sp":
                    parent_circuit_entry_id = str(entry.get("parent_circuit_entry_id") or "").strip()
                    if parent_circuit_entry_id and parent_circuit_entry_id in self._window_registry["circuit"]:
                        nested_plot_entries.setdefault(parent_circuit_entry_id, []).append((kind, entry_id, title))
                        continue
                flat_entries.append((kind, entry_id, title))
        flat_entries.sort(key=lambda row: row[2].lower())
        if not flat_entries:
            project_item.addChild(QTreeWidgetItem(["(none)"]))
        else:
            circuit_items: dict[str, QTreeWidgetItem] = {}
            for kind, entry_id, title in flat_entries:
                item = self._add_tree_window_item(project_item, kind, entry_id, title)
                if kind == "circuit":
                    circuit_items[entry_id] = item

            for circuit_entry_id, entries in nested_plot_entries.items():
                parent_item = circuit_items.get(circuit_entry_id)
                if parent_item is None:
                    continue
                for kind, entry_id, title in sorted(entries, key=lambda row: row[2].lower()):
                    self._add_tree_window_item(parent_item, kind, entry_id, title)
                parent_item.setExpanded(True)

        project_item.setExpanded(True)

    def _on_project_tree_context_menu(self, pos) -> None:
        item = self._project_tree.itemAt(pos)
        if item is None:
            return

        payload = item.data(0, Qt.UserRole)
        if not isinstance(payload, dict):
            return

        action = str(payload.get("action", ""))
        if action not in {"window", "output-file"}:
            return

        self._project_tree.setCurrentItem(item)

        menu = QMenu(self)
        open_action = menu.addAction("Open")
        duplicate_action = None
        if self._tree_payload_can_duplicate(payload):
            duplicate_action = menu.addAction("Duplicate")
        close_action = menu.addAction("Close")
        minimize_action = menu.addAction("Minimize")
        menu.addSeparator()
        delete_action = menu.addAction("Delete")

        is_open = self._tree_payload_is_open(payload)
        close_action.setEnabled(is_open)
        minimize_action.setEnabled(is_open)

        chosen = menu.exec(self._project_tree.viewport().mapToGlobal(pos))
        if chosen is open_action:
            self._open_tree_payload(payload)
            return
        if duplicate_action is not None and chosen is duplicate_action:
            self._duplicate_tree_payload(payload)
            return
        if chosen is close_action:
            self._close_tree_payload(payload)
            return
        if chosen is minimize_action:
            self._minimize_tree_payload(payload)
            return
        if chosen is delete_action:
            self._confirm_and_delete_tree_payload(item, payload)

    def _open_tree_payload(self, payload: dict) -> None:
        action = str(payload.get("action", ""))
        if action == "window":
            self._activate_or_open_window_entry(
                str(payload.get("kind", "")),
                str(payload.get("id", "")),
            )
            return
        if action == "output-file":
            self._open_output_file_from_tree(
                str(payload.get("file", "")),
                str(payload.get("output", "")),
            )
            return
        if action == "transient-window":
            self._activate_transient_window(
                str(payload.get("id", "")),
                int(payload.get("index", 0)),
            )

    def _activate_transient_window(self, circuit_entry_id: str, index: int) -> None:
        circuit_widget = self._open_windows.get(("circuit", circuit_entry_id))
        if circuit_widget is None or not hasattr(circuit_widget, "get_open_transient_windows"):
            return
        try:
            windows = list(circuit_widget.get_open_transient_windows())
        except Exception:
            return
        if index < 0 or index >= len(windows):
            return
        target = windows[index]
        self._windows.set_active_widget(target)
        target.raise_()
        target.activateWindow()

    def _tree_payload_can_duplicate(self, payload: dict) -> bool:
        action = str(payload.get("action", ""))
        if action == "window":
            return str(payload.get("kind", "")) in {"sp", "tdr", "circuit"}
        if action == "output-file":
            return str(payload.get("output", "")) == "eye"
        return False

    def _duplicate_tree_payload(self, payload: dict) -> bool:
        action = str(payload.get("action", ""))
        if action == "window":
            return self._duplicate_tree_window_entry(
                str(payload.get("kind", "")),
                str(payload.get("id", "")),
            )
        if action == "output-file":
            return self._duplicate_tree_output_file(
                str(payload.get("file", "")),
                str(payload.get("output", "")),
            )
        return False

    def _duplicate_tree_window_entry(self, kind: str, entry_id: str) -> bool:
        entry = self._window_registry.get(kind, {}).get(entry_id)
        if entry is None:
            return False

        source_widget = self._open_windows.get((kind, entry_id))
        state = {}
        if source_widget is not None:
            state = deepcopy(self._snapshot_widget_state(source_widget))
        elif isinstance(entry.get("state"), dict):
            state = deepcopy(entry["state"])

        required_paths = self._touchstone_paths_for_window_state(kind, state)
        if not self._ensure_touchstone_files_loaded(required_paths, kind=kind):
            return False

        source_size = self._window_size_for_duplicate(entry, source_widget)

        if kind == "sp":
            return self._open_duplicated_plot_window(
                state,
                source_size,
                parent_circuit_entry_id=str(entry.get("parent_circuit_entry_id") or "").strip() or None,
            )
        if kind == "tdr":
            return self._open_duplicated_tdr_window(state, source_size)
        if kind == "circuit":
            return self._open_duplicated_circuit_window(state, source_size)
        return False

    def _touchstone_paths_for_window_state(self, kind: str, state: dict) -> list[str]:
        paths: list[str] = []
        seen: set[str] = set()

        if kind in {"sp", "tdr"}:
            for file_entry in state.get("files", []):
                if not isinstance(file_entry, dict):
                    continue
                raw_path = str(file_entry.get("file_path") or "").strip()
                if not raw_path:
                    continue
                try:
                    normalized = str(Path(raw_path).resolve())
                except OSError:
                    normalized = raw_path
                if normalized in seen:
                    continue
                seen.add(normalized)
                paths.append(normalized)
            return paths

        if kind == "circuit":
            for instance in state.get("instances", []):
                if not isinstance(instance, dict):
                    continue
                if str(instance.get("block_kind", "touchstone")) != "touchstone":
                    continue
                raw_path = str(instance.get("source_file_id") or "").strip()
                if not raw_path:
                    continue
                try:
                    normalized = str(Path(raw_path).resolve())
                except OSError:
                    normalized = raw_path
                if normalized in seen:
                    continue
                seen.add(normalized)
                paths.append(normalized)

        return paths

    def _ensure_touchstone_files_loaded(self, paths: list[str], *, kind: str) -> bool:
        requested = [path for path in paths if path]
        if not requested:
            return True

        _added, errors = self._state.load_files(requested)
        loaded_paths = {str(item.path.resolve()) for item in self._state.get_loaded_files()}
        missing = [path for path in requested if path not in loaded_paths]
        if not errors and not missing:
            return True

        details = errors[:]
        details.extend(f"{Path(path).name}: file not available" for path in missing)
        kind_label = {
            "sp": "plot",
            "tdr": "TDR plot",
            "circuit": "circuit",
        }.get(kind, "item")
        QMessageBox.warning(
            self,
            "Duplicate item",
            f"Could not duplicate this {kind_label} because some Touchstone files could not be loaded:\n\n" + "\n".join(details),
        )
        return False

    def _window_size_for_duplicate(self, entry: dict, widget: object | None) -> list[int] | None:
        if widget is not None and hasattr(widget, "size"):
            size = widget.size()
            if size.width() > 0 and size.height() > 0:
                return [size.width(), size.height()]
        window_size = entry.get("window_size")
        if isinstance(window_size, list) and len(window_size) == 2:
            return [int(window_size[0]), int(window_size[1])]
        return None

    def _copy_label(self, label: str, fallback: str) -> str:
        base = str(label).strip() or fallback
        return f"{base} Copy"

    def _rename_duplicated_widget(self, widget: object) -> None:
        if isinstance(widget, PlotWindow):
            copy_name = self._copy_label(getattr(widget, "_plot_name", ""), f"Plot #{widget.window_number}")
            widget._plot_name = copy_name
            widget._plot_name_edit.blockSignals(True)
            widget._plot_name_edit.setText(copy_name)
            widget._plot_name_edit.blockSignals(False)
            widget._refresh_window_title()
            return

        if isinstance(widget, TdrWindow):
            copy_name = self._copy_label(getattr(widget, "_tdr_name", ""), f"TDR #{widget.window_number}")
            widget._tdr_name = copy_name
            widget._tdr_name_edit.blockSignals(True)
            widget._tdr_name_edit.setText(copy_name)
            widget._tdr_name_edit.blockSignals(False)
            widget._refresh_window_title()
            return

        if isinstance(widget, CircuitWindow):
            copy_name = self._copy_label(widget.circuit_display_name(), f"Circuit #{widget.window_number}")
            widget._circuit_name = copy_name
            widget._circuit_name_edit.blockSignals(True)
            widget._circuit_name_edit.setText(copy_name)
            widget._circuit_name_edit.blockSignals(False)
            widget._refresh_window_title()
            widget._sync_eye_window_titles()

    def _open_duplicated_plot_window(
        self,
        state: dict,
        window_size: list[int] | None,
        *,
        parent_circuit_entry_id: str | None = None,
    ) -> bool:
        self._plot_counter += 1
        plot_win = PlotWindow(self._state, window_number=self._plot_counter)
        plot_win.apply_project_state(state)
        self._rename_duplicated_widget(plot_win)
        plot_win.project_modified.connect(self._mark_project_dirty)

        if isinstance(window_size, list) and len(window_size) == 2:
            plot_win.resize(int(window_size[0]), int(window_size[1]))
        else:
            plot_win.resize(*default_child_window_size("plots"))
        entry_id = self._register_window_entry(
            "sp",
            plot_win,
            parent_circuit_entry_id=parent_circuit_entry_id,
        )
        self._bind_open_window("sp", entry_id, plot_win)
        self._windows.present(plot_win)
        self._mark_project_dirty()
        self._refresh_project_tree()
        return True

    def _open_duplicated_tdr_window(self, state: dict, window_size: list[int] | None) -> bool:
        self._tdr_counter += 1
        tdr_win = TdrWindow(self._state, window_number=self._tdr_counter)
        tdr_win.apply_project_state(state)
        self._rename_duplicated_widget(tdr_win)
        tdr_win.project_modified.connect(self._mark_project_dirty)

        if isinstance(window_size, list) and len(window_size) == 2:
            tdr_win.resize(int(window_size[0]), int(window_size[1]))
        else:
            tdr_win.resize(*default_child_window_size("tdr"))
        entry_id = self._register_window_entry("tdr", tdr_win)
        self._bind_open_window("tdr", entry_id, tdr_win)
        self._windows.present(tdr_win)
        self._mark_project_dirty()
        self._refresh_project_tree()
        return True

    def _open_duplicated_circuit_window(self, state: dict, window_size: list[int] | None) -> bool:
        circuit_number = self._next_available_circuit_number()
        self._circuit_counter = max(self._circuit_counter, circuit_number)
        circuit_win = CircuitWindow(self._state, window_number=circuit_number)
        circuit_win.apply_project_state(state)
        self._rename_duplicated_widget(circuit_win)
        circuit_win.project_modified.connect(self._mark_project_dirty)

        if isinstance(window_size, list) and len(window_size) == 2:
            circuit_win.resize(int(window_size[0]), int(window_size[1]))
        else:
            circuit_win.resize(*default_child_window_size("circuits"))
        entry_id = self._register_window_entry("circuit", circuit_win)
        self._bind_open_window("circuit", entry_id, circuit_win)
        self._windows.present(circuit_win)
        self._mark_project_dirty()
        self._refresh_project_tree()
        return True

    def _duplicate_tree_output_file(self, file_name: str, output_kind: str) -> bool:
        if output_kind != "eye" or not file_name:
            return False
        if self._project_data_dir is None:
            QMessageBox.information(self, "Duplicate item", "Project data folder is not available.")
            return False

        output_path = self._project_data_dir / file_name
        if not output_path.exists():
            QMessageBox.warning(self, "Duplicate item", f"Output file not found:\n{output_path}")
            return False

        source_window = next(iter(self._find_output_eye_windows(file_name, output_kind)), None)
        if isinstance(source_window, EyeDiagramWindow):
            duplicated = self._show_eye_diagram_window(
                source_window._result,
                title=self._copy_label(source_window.windowTitle(), output_path.stem),
                eye_span_ui=int(source_window._eye_span_ui),
                render_mode=str(source_window._render_mode),
                quality_preset=str(source_window._quality_preset),
                statistical_enabled=bool(source_window._statistical_enabled),
                noise_rms_mv=float(source_window._noise_rms_mv),
                jitter_rms_ps=float(source_window._jitter_rms_ps),
                output_file_name=file_name,
            )
            return duplicated is not None

        duplicated = self._open_eye_binary_file(
            output_path,
            title_override=self._copy_label(self._tree_display_name(file_name), output_path.stem),
        )
        return duplicated is not None

    def _tree_payload_is_open(self, payload: dict) -> bool:
        action = str(payload.get("action", ""))
        if action == "window":
            return self._find_open_window_subwindow(
                str(payload.get("kind", "")),
                str(payload.get("id", "")),
            ) is not None
        if action == "output-file":
            return bool(
                self._find_output_targets(
                    str(payload.get("file", "")),
                    str(payload.get("output", "")),
                )
            )
        return False

    def _find_open_window_subwindow(self, kind: str, entry_id: str):
        return self._open_windows.get((kind, entry_id))

    def _output_target_matches(self, target: object, file_name: str, output_kind: str) -> bool:
        if target is None:
            return False
        return (
            str(getattr(target, "property", lambda *_: None)("tree_output_file") or "") == file_name
            and str(getattr(target, "property", lambda *_: None)("tree_output_kind") or "") == output_kind
        )

    def _find_output_targets(self, file_name: str, output_kind: str) -> list[QWidget]:
        if not file_name or not output_kind:
            return []

        matches: list[QWidget] = []
        seen_ids: set[int] = set()
        for widget in self._windows.list_widgets():
            if self._output_target_matches(widget, file_name, output_kind):
                matches.append(widget)
                seen_ids.add(id(widget))

        for widget in self.findChildren(QWidget):
            if id(widget) in seen_ids:
                continue
            if self._output_target_matches(widget, file_name, output_kind):
                matches.append(widget)
                seen_ids.add(id(widget))

        return matches

    def _find_output_eye_windows(self, file_name: str, output_kind: str) -> list[EyeDiagramWindow]:
        if not file_name or not output_kind:
            return []

        matches: list[EyeDiagramWindow] = []
        seen_ids: set[int] = set()
        for widget in self._windows.widgets_of_type(EyeDiagramWindow):
            if self._output_target_matches(widget, file_name, output_kind):
                matches.append(widget)
                seen_ids.add(id(widget))

        for widget in self.findChildren(EyeDiagramWindow):
            if id(widget) in seen_ids:
                continue
            if self._output_target_matches(widget, file_name, output_kind):
                matches.append(widget)
                seen_ids.add(id(widget))

        return matches

    def _close_tree_payload(self, payload: dict) -> bool:
        action = str(payload.get("action", ""))
        if action == "window":
            sub = self._find_open_window_subwindow(
                str(payload.get("kind", "")),
                str(payload.get("id", "")),
            )
            if sub is None:
                return False
            return bool(sub.close())

        if action == "output-file":
            closed_any = False
            for target in self._find_output_targets(
                str(payload.get("file", "")),
                str(payload.get("output", "")),
            ):
                target.close()
                closed_any = True
            return closed_any

        return False

    def _minimize_tree_payload(self, payload: dict) -> bool:
        action = str(payload.get("action", ""))
        if action == "window":
            sub = self._find_open_window_subwindow(
                str(payload.get("kind", "")),
                str(payload.get("id", "")),
            )
            if sub is None:
                return False
            sub.showMinimized()
            return True

        if action == "output-file":
            minimized_any = False
            for target in self._find_output_targets(
                str(payload.get("file", "")),
                str(payload.get("output", "")),
            ):
                target.showMinimized()
                minimized_any = True
            return minimized_any

        return False

    def _confirm_and_delete_tree_payload(self, item: QTreeWidgetItem, payload: dict) -> None:
        label = item.text(0).strip() or "selected item"
        action = str(payload.get("action", ""))
        detail = "This action cannot be undone."
        if action == "window":
            detail = "This will remove the item from the project tree."
            if str(payload.get("kind", "")) == "circuit":
                detail = "This will remove the circuit from the project tree and delete any saved Eye and S-parameter output files associated with it."
        elif action == "output-file":
            detail = "This will delete the saved output file from disk and remove it from the project tree."

        result = QMessageBox.question(
            self,
            "Delete item",
            f"Delete '{label}'?\n\n{detail}",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No,
        )
        if result != QMessageBox.Yes:
            return

        self._delete_tree_payload(payload)

    def _delete_tree_payload(self, payload: dict) -> bool:
        action = str(payload.get("action", ""))
        if action == "window":
            return self._delete_tree_window_entry(
                str(payload.get("kind", "")),
                str(payload.get("id", "")),
            )
        if action == "output-file":
            return self._delete_tree_output_file(
                str(payload.get("kind", "")),
                str(payload.get("id", "")),
                str(payload.get("file", "")),
                str(payload.get("output", "")),
            )
        return False

    def _delete_tree_window_entry(self, kind: str, entry_id: str) -> bool:
        entry = self._window_registry.get(kind, {}).get(entry_id)
        if entry is None:
            return False

        sub = self._find_open_window_subwindow(kind, entry_id)
        if sub is not None and not sub.close():
            return False

        if kind == "circuit":
            eye_file = str(entry.get("eye_file") or "").strip()
            sparam_file = str(entry.get("sparam_file") or "").strip()
            sparam_plot_file = str(entry.get("sparam_plot_file") or "").strip()
            transient_file = str(entry.get("transient_file") or "").strip()

            related_plot_ids = [
                plot_entry_id
                for plot_entry_id, plot_entry in self._window_registry["sp"].items()
                if str(plot_entry.get("parent_circuit_entry_id") or "") == entry_id
            ]
            for plot_entry_id in related_plot_ids:
                self._delete_tree_window_entry("sp", plot_entry_id)

            if eye_file:
                self._close_output_targets(eye_file, "eye")
                self._delete_runtime_file(eye_file)
            if sparam_plot_file:
                self._close_output_targets(sparam_plot_file, "sparam-plot")
                self._delete_runtime_file(sparam_plot_file)
            if transient_file:
                self._close_output_targets(transient_file, "transient")
                self._delete_runtime_file(transient_file)
            if sparam_file:
                if self._project_data_dir is not None:
                    self._state.unload_file(str((self._project_data_dir / sparam_file).resolve()))
                self._delete_runtime_file(sparam_file)

        self._open_windows.pop((kind, entry_id), None)
        self._window_registry.get(kind, {}).pop(entry_id, None)
        self._mark_project_dirty()
        self._refresh_project_tree()
        return True

    def _close_output_targets(self, file_name: str, output_kind: str) -> bool:
        closed_any = False
        for target in self._find_output_targets(file_name, output_kind):
            target.close()
            closed_any = True
        return closed_any

    def _delete_tree_output_file(self, kind: str, entry_id: str, file_name: str, output_kind: str) -> bool:
        entry = self._window_registry.get(kind, {}).get(entry_id)
        if entry is None or not file_name:
            return False

        self._close_output_targets(file_name, output_kind)
        self._delete_runtime_file(file_name)

        if output_kind == "eye":
            entry["eye_file"] = None
        elif output_kind == "sparam-plot":
            entry["sparam_plot_file"] = None
        elif output_kind == "transient":
            entry["transient_file"] = None

        if not any(str(entry.get(key) or "").strip() for key in ("eye_file", "sparam_file", "sparam_plot_file", "transient_file")):
            entry["output_kind"] = None

        self._mark_project_dirty()
        self._refresh_project_tree()
        return True

    def _delete_runtime_file(self, file_name: str) -> bool:
        if not file_name:
            return False

        data_dir = self._ensure_data_dir_for_runtime_exports()
        if data_dir is None:
            return False

        file_path = data_dir / file_name
        if not file_path.exists():
            return False

        try:
            file_path.unlink()
        except OSError:
            return False
        return True

    def _on_project_tree_item_clicked(self, item: QTreeWidgetItem, _column: int) -> None:
        payload = item.data(0, Qt.UserRole)
        if not isinstance(payload, dict):
            return
        action = str(payload.get("action", ""))
        if action == "window":
            self._activate_or_open_window_entry(
                str(payload.get("kind", "")),
                str(payload.get("id", "")),
            )
            return
        if action == "output-file":
            self._open_output_file_from_tree(
                str(payload.get("file", "")),
                str(payload.get("output", "")),
            )
            return
        if action == "transient-window":
            self._activate_transient_window(
                str(payload.get("id", "")),
                int(payload.get("index", 0)),
            )

    def _activate_or_open_window_entry(self, kind: str, entry_id: str) -> None:
        key = (kind, entry_id)
        open_widget = self._open_windows.get(key)
        if open_widget is not None:
            self._windows.set_active_widget(open_widget)
            return

        entry = self._window_registry.get(kind, {}).get(entry_id)
        if entry is None:
            return
        state = entry.get("state") if isinstance(entry.get("state"), dict) else {}
        win_number = int(entry.get("window_number", 0))
        window_size = entry.get("window_size")

        if kind == "sp":
            self._plot_counter = max(self._plot_counter, win_number)
            widget = PlotWindow(self._state, window_number=max(1, win_number))
            if state:
                widget.apply_project_state(state)
            widget.project_modified.connect(self._mark_project_dirty)
            if isinstance(window_size, list) and len(window_size) == 2:
                widget.resize(int(window_size[0]), int(window_size[1]))
            else:
                widget.resize(*default_child_window_size("plots"))
            self._bind_open_window("sp", entry_id, widget)
            self._windows.present(widget)
            self._refresh_project_tree()
            return

        if kind == "tdr":
            self._tdr_counter = max(self._tdr_counter, win_number)
            widget = TdrWindow(self._state, window_number=max(1, win_number))
            if state:
                widget.apply_project_state(state)
            widget.project_modified.connect(self._mark_project_dirty)
            if isinstance(window_size, list) and len(window_size) == 2:
                widget.resize(int(window_size[0]), int(window_size[1]))
            else:
                widget.resize(*default_child_window_size("tdr"))
            self._bind_open_window("tdr", entry_id, widget)
            self._windows.present(widget)
            self._refresh_project_tree()
            return

        if kind == "circuit":
            self._circuit_counter = max(self._circuit_counter, win_number)
            widget = CircuitWindow(self._state, window_number=max(1, win_number))
            if state:
                widget.apply_project_state(state)
            widget.project_modified.connect(self._mark_project_dirty)
            if isinstance(window_size, list) and len(window_size) == 2:
                widget.resize(int(window_size[0]), int(window_size[1]))
            else:
                widget.resize(*default_child_window_size("circuits"))
            self._bind_open_window("circuit", entry_id, widget)
            self._windows.present(widget)
            self._refresh_project_tree()

    def _open_output_file_from_tree(self, file_name: str, output_kind: str) -> None:
        if not file_name:
            return
        if self._project_data_dir is None:
            QMessageBox.information(self, "Output file", "Project data folder is not available.")
            return
        output_path = self._project_data_dir / file_name
        if not output_path.exists():
            QMessageBox.warning(self, "Output file", f"Output file not found:\n{output_path}")
            return

        if output_kind == "eye":
            self._open_eye_binary_file(output_path)
            return

        if output_kind == "transient":
            self._open_transient_binary_file(output_path)
            return

        if output_kind == "sparam-plot":
            pixmap = QPixmap(str(output_path))
            if pixmap.isNull():
                QMessageBox.warning(self, "Output file", f"Could not load plot image:\n{output_path}")
                return
            label = QLabel()
            label.setPixmap(pixmap)
            label.setAlignment(Qt.AlignCenter)
            scroll = QScrollArea()
            scroll.setWidget(label)
            scroll.setWidgetResizable(True)
            scroll.setWindowTitle(file_name)
            scroll.setProperty("tree_output_file", file_name)
            scroll.setProperty("tree_output_kind", output_kind)
            scroll.resize(*default_child_window_size("plots"))
            self._windows.present(scroll)

    def _load_eye_binary_snapshot(self, file_path: Path) -> tuple[ChannelSimResult, dict[str, object]] | None:
        try:
            with np.load(file_path, allow_pickle=False) as data:
                time_s = np.array(data["time_s"], dtype=float)
                waveform_v = np.array(data["waveform_v"], dtype=float)
                ui_s = float(np.asarray(data["ui_s"]).reshape(-1)[0])
                is_diff = bool(int(np.asarray(data["is_differential"]).reshape(-1)[0]))
                spec_json = str(np.asarray(data["driver_spec_json"]).reshape(-1)[0])
                spec_payload = json.loads(spec_json)
                eye_span_ui = int(np.asarray(data["eye_span_ui"]).reshape(-1)[0])
                render_mode = str(np.asarray(data["render_mode"]).reshape(-1)[0])
                quality_preset = str(np.asarray(data["quality_preset"]).reshape(-1)[0])
                stat_enabled = bool(int(np.asarray(data["stat_enabled"]).reshape(-1)[0]))
                noise_rms_mv = float(np.asarray(data["noise_rms_mv"]).reshape(-1)[0])
                jitter_rms_ps = float(np.asarray(data["jitter_rms_ps"]).reshape(-1)[0])
        except Exception as exc:
            QMessageBox.warning(self, "Eye file", f"Could not load eye binary:\n{exc}")
            return None

        result = ChannelSimResult(
            time_s=time_s,
            waveform_v=waveform_v,
            ui_s=ui_s,
            driver_spec=DriverSpec.from_dict(spec_payload),
            is_differential=is_diff,
        )
        return result, {
            "eye_span_ui": eye_span_ui,
            "render_mode": render_mode,
            "quality_preset": quality_preset,
            "statistical_enabled": stat_enabled,
            "noise_rms_mv": noise_rms_mv,
            "jitter_rms_ps": jitter_rms_ps,
        }

    def _show_eye_diagram_window(
        self,
        result: ChannelSimResult,
        *,
        title: str,
        eye_span_ui: int,
        render_mode: str,
        quality_preset: str,
        statistical_enabled: bool,
        noise_rms_mv: float,
        jitter_rms_ps: float,
        output_file_name: str | None = None,
    ) -> EyeDiagramWindow:
        eye_win = EyeDiagramWindow(
            result,
            title=title,
            initial_span_ui=eye_span_ui,
            initial_render_mode=render_mode,
            initial_quality_preset=quality_preset,
            statistical_enabled=statistical_enabled,
            noise_rms_mv=noise_rms_mv,
            jitter_rms_ps=jitter_rms_ps,
        )
        if output_file_name:
            eye_win.setProperty("tree_output_file", output_file_name)
        eye_win.setProperty("tree_output_kind", "eye")
        eye_win.resize(*default_child_window_size("eye"))
        self._windows.present(eye_win)
        return eye_win

    def _open_eye_binary_file(self, file_path: Path, *, title_override: str | None = None) -> EyeDiagramWindow | None:
        snapshot = self._load_eye_binary_snapshot(file_path)
        if snapshot is None:
            return None

        result, settings = snapshot
        return self._show_eye_diagram_window(
            result,
            title=title_override or file_path.stem,
            eye_span_ui=int(settings["eye_span_ui"]),
            render_mode=str(settings["render_mode"]),
            quality_preset=str(settings["quality_preset"]),
            statistical_enabled=bool(settings["statistical_enabled"]),
            noise_rms_mv=float(settings["noise_rms_mv"]),
            jitter_rms_ps=float(settings["jitter_rms_ps"]),
            output_file_name=file_path.name,
        )

    # ── Transient binary persistence ─────────────────────────────────────

    def _build_transient_binary_payload(self, result: TransientSimResult) -> bytes:
        traces = result.traces
        n_traces = len(traces)
        n_samples = int(result.time_s.size)
        if n_traces > 0:
            waveforms = np.stack(
                [np.asarray(tr.waveform_v, dtype=float) for tr in traces], axis=0
            )
        else:
            waveforms = np.zeros((0, n_samples), dtype=float)

        trace_metadata = [
            {
                "output_instance_id": tr.output_instance_id,
                "output_port_number": int(tr.output_port_number),
                "label": tr.label,
            }
            for tr in traces
        ]
        buf = BytesIO()
        np.savez_compressed(
            buf,
            time_s=np.asarray(result.time_s, dtype=float),
            waveforms_v=waveforms,
            trace_metadata_json=np.asarray([json.dumps(trace_metadata)]),
            source_spec_json=np.asarray([json.dumps(result.source_spec.to_dict())]),
            warnings_json=np.asarray([json.dumps(list(result.warnings))]),
        )
        return buf.getvalue()

    def _load_transient_binary_snapshot(self, file_path: Path) -> TransientSimResult | None:
        try:
            with np.load(file_path, allow_pickle=False) as data:
                time_s = np.array(data["time_s"], dtype=float)
                waveforms = np.array(data["waveforms_v"], dtype=float)
                trace_metadata = json.loads(
                    str(np.asarray(data["trace_metadata_json"]).reshape(-1)[0])
                )
                source_spec_payload = json.loads(
                    str(np.asarray(data["source_spec_json"]).reshape(-1)[0])
                )
                warnings_list = json.loads(
                    str(np.asarray(data["warnings_json"]).reshape(-1)[0])
                )
        except Exception as exc:
            QMessageBox.warning(self, "Transient file", f"Could not load transient binary:\n{exc}")
            return None

        traces: list[TransientTrace] = []
        for index, meta in enumerate(trace_metadata):
            if not isinstance(meta, dict) or index >= waveforms.shape[0]:
                continue
            traces.append(
                TransientTrace(
                    output_instance_id=str(meta.get("output_instance_id", "")),
                    output_port_number=int(meta.get("output_port_number", 0)),
                    label=str(meta.get("label", f"Trace {index + 1}")),
                    waveform_v=np.asarray(waveforms[index], dtype=float),
                )
            )
        return TransientSimResult(
            time_s=time_s,
            traces=tuple(traces),
            source_spec=TransientSourceSpec.from_dict(source_spec_payload),
            warnings=tuple(str(w) for w in warnings_list),
        )

    def _open_transient_binary_file(
        self, file_path: Path, *, title_override: str | None = None
    ) -> TransientResultWindow | None:
        result = self._load_transient_binary_snapshot(file_path)
        if result is None:
            return None
        title = title_override or file_path.stem
        existing = self._find_transient_window_for_file(file_path.name)
        if existing is not None:
            existing.update_result(result)
            existing.setWindowTitle(title)
            self._windows.set_active_widget(existing)
            existing.raise_()
            existing.activateWindow()
            return existing
        win = TransientResultWindow(result, title=title)
        win.setProperty("tree_output_file", file_path.name)
        win.setProperty("tree_output_kind", "transient")
        win.resize(*default_child_window_size("transient"))
        self._windows.present(win)
        return win

    def _find_transient_window_for_file(self, file_name: str) -> TransientResultWindow | None:
        if not file_name:
            return None
        for widget in self._windows.widgets_of_type(TransientResultWindow):
            if str(widget.property("tree_output_file") or "") == file_name:
                return widget
        return None

    def _on_circuit_transient_result_generated(self, entry_id: str, payload: object) -> None:
        entry = self._window_registry.get("circuit", {}).get(entry_id)
        if entry is None or not isinstance(payload, dict):
            return
        result = payload.get("result")
        if not isinstance(result, TransientSimResult):
            return

        data_dir = self._ensure_data_dir_for_runtime_exports()
        if data_dir is None:
            return

        circuit_name = str(payload.get("circuit_name") or entry.get("title") or f"Circuit_{entry_id}")
        transient_file_name = (
            self._sanitize_file_stem(circuit_name, f"Circuit_{entry_id}") + "_Transient.tran"
        )
        transient_path = data_dir / transient_file_name

        try:
            transient_path.write_bytes(self._build_transient_binary_payload(result))
        except Exception:
            return

        old_transient_file = str(entry.get("transient_file") or "").strip()
        entry["transient_file"] = transient_file_name
        if old_transient_file and old_transient_file != transient_file_name:
            old_path = data_dir / old_transient_file
            if old_path.exists():
                try:
                    old_path.unlink()
                except OSError:
                    pass
        # Tag any live TransientResultWindow opened from this circuit so that
        # closing/reopening from the tree reuses it instead of spawning a new one.
        for inner in self._windows.widgets_of_type(TransientResultWindow):
            inner.setProperty("tree_output_file", transient_file_name)
            inner.setProperty("tree_output_kind", "transient")
        self._mark_project_dirty()
        self._refresh_project_tree()

    # ── File loading ──────────────────────────────────────────────────────

    def _settings(self) -> QSettings:
        return QSettings("SPUtility", "SPUtility")

    def _load_recent_entries(self) -> None:
        settings = self._settings()
        recent_projects = settings.value("recent_projects", [], type=list)
        recent_sparams = settings.value("recent_sparams", [], type=list)

        self._recent_projects = [
            str(p) for p in recent_projects if isinstance(p, str) and p.strip()
        ][:_RECENT_ENTRIES_MAX]
        self._recent_sparams = [
            str(p) for p in recent_sparams if isinstance(p, str) and p.strip()
        ][:_RECENT_ENTRIES_MAX]

    def _save_recent_entries(self) -> None:
        settings = self._settings()
        settings.setValue("recent_projects", self._recent_projects)
        settings.setValue("recent_sparams", self._recent_sparams)

    def _normalize_path(self, file_path: str) -> str:
        try:
            return str(Path(file_path).expanduser().resolve())
        except (OSError, RuntimeError):
            return str(file_path)

    def _push_recent_project(self, file_path: str) -> None:
        normalized = self._normalize_path(file_path)
        self._recent_projects = [p for p in self._recent_projects if p != normalized]
        self._recent_projects.insert(0, normalized)
        self._recent_projects = self._recent_projects[:_RECENT_ENTRIES_MAX]
        self._save_recent_entries()
        self._rebuild_recent_file_menus()

    def _push_recent_sparam(self, file_path: str) -> None:
        normalized = self._normalize_path(file_path)
        self._recent_sparams = [p for p in self._recent_sparams if p != normalized]
        self._recent_sparams.insert(0, normalized)
        self._recent_sparams = self._recent_sparams[:_RECENT_ENTRIES_MAX]
        self._save_recent_entries()
        self._rebuild_recent_file_menus()

    def _clear_recent_projects(self) -> None:
        self._recent_projects = []
        self._save_recent_entries()
        self._rebuild_recent_file_menus()

    def _clear_recent_sparams(self) -> None:
        self._recent_sparams = []
        self._save_recent_entries()
        self._rebuild_recent_file_menus()

    def _rebuild_recent_file_menus(self) -> None:
        self._recent_projects_menu.clear()
        if not self._recent_projects:
            action = self._recent_projects_menu.addAction("No recent projects")
            action.setEnabled(False)
        else:
            for index, path in enumerate(self._recent_projects, start=1):
                label = f"{index}. {Path(path).name}"
                action = self._recent_projects_menu.addAction(label)
                action.setToolTip(path)
                action.triggered.connect(
                    lambda checked=False, p=path: self._open_recent_project(p)
                )
            self._recent_projects_menu.addSeparator()
            self._recent_projects_menu.addAction("Clear project history", self._clear_recent_projects)

        self._recent_sparams_menu.clear()
        if not self._recent_sparams:
            action = self._recent_sparams_menu.addAction("No recent S-parameter files")
            action.setEnabled(False)
        else:
            touchstone_icon = self._icon_for_kind("touchstone-file")
            for index, path in enumerate(self._recent_sparams, start=1):
                label = f"{index}. {Path(path).name}"
                action = self._recent_sparams_menu.addAction(label)
                action.setIcon(touchstone_icon)
                action.setToolTip(path)
                action.triggered.connect(
                    lambda checked=False, p=path: self._open_recent_sparam(p)
                )
            self._recent_sparams_menu.addSeparator()
            self._recent_sparams_menu.addAction("Clear S-parameter history", self._clear_recent_sparams)

    def _open_recent_project(self, file_path: str) -> None:
        if not Path(file_path).exists():
            self._recent_projects = [p for p in self._recent_projects if p != file_path]
            self._save_recent_entries()
            self._rebuild_recent_file_menus()
            QMessageBox.warning(self, "Missing file", f"Project file not found:\n{file_path}")
            return
        self._load_project_from_path(file_path)

    def _open_recent_sparam(self, file_path: str) -> None:
        if not Path(file_path).exists():
            self._recent_sparams = [p for p in self._recent_sparams if p != file_path]
            self._save_recent_entries()
            self._rebuild_recent_file_menus()
            QMessageBox.warning(self, "Missing file", f"S-parameter file not found:\n{file_path}")
            return
        self._load_touchstone_files([file_path])

    def _load_touchstone_files(self, files: list[str]) -> None:
        if not files:
            return

        added_count, errors = self._state.load_files(files)

        if errors:
            QMessageBox.warning(self, "Load errors", "\n".join(errors))

        if added_count == 0 and not errors:
            QMessageBox.information(
                self, "No new files", "The selected files are already loaded."
            )
        elif added_count > 0:
            self._mark_project_dirty()

        for path in files:
            if Path(path).exists():
                self._push_recent_sparam(path)

    def _open_files(self) -> None:
        files, _ = QFileDialog.getOpenFileNames(
            self,
            "Open Touchstone file",
            "",
            "Touchstone (*.s1p *.s2p *.s3p *.s4p *.s5p *.s6p *.s7p *.s8p *.s9p"
            " *.s10p *.s12p *.s16p *.ts);;All files (*)",
        )
        if not files:
            return

        self._load_touchstone_files(files)

    def _build_project_payload(self) -> dict:
        loaded_files = self._state.get_loaded_files()
        files_data = [
            {
                "file_path": str(loaded.path),
                "file_name": loaded.display_name,
            }
            for loaded in loaded_files
        ]

        for (kind, entry_id), widget in list(self._open_windows.items()):
            self._sync_window_entry(kind, entry_id, widget)

        plot_windows: list[dict] = []
        tdr_windows: list[dict] = []
        circuit_windows: list[dict] = []
        registry_payload: dict[str, list[dict]] = {"sp": [], "tdr": [], "circuit": []}

        for kind in ("sp", "tdr", "circuit"):
            for entry in self._window_registry[kind].values():
                out = dict(entry)
                registry_payload[kind].append(out)
                if not bool(entry.get("is_open", False)):
                    continue
                state = dict(entry.get("state", {}))
                window_size = entry.get("window_size")
                if isinstance(window_size, list) and len(window_size) == 2:
                    state["window_size"] = [int(window_size[0]), int(window_size[1])]
                if kind == "sp":
                    plot_windows.append(state)
                elif kind == "tdr":
                    tdr_windows.append(state)
                else:
                    circuit_windows.append(state)

        return {
            "app_name": pkg.__app_name__,
            "app_version": pkg.__version__,
            "saved_at": datetime.now().isoformat(timespec="seconds"),
            "loaded_files": files_data,
            "plots": plot_windows,
            "tdr_plots": tdr_windows,
            "circuits": circuit_windows,
            "window_registry": registry_payload,
        }

    def _resolve_project_reference_path(self, project_dir: Path, raw_path: str) -> str:
        text = str(raw_path).strip()
        if not text:
            return text
        path = Path(text)
        if not path.is_absolute():
            path = project_dir / path
        return self._normalize_path(str(path))

    def _rewrite_window_state_touchstone_paths(self, kind: str, state: dict, transform) -> dict:
        updated = deepcopy(state)

        if kind in {"sp", "tdr"}:
            rewritten_files: list[dict] = []
            for file_entry in updated.get("files", []):
                if not isinstance(file_entry, dict):
                    continue
                rewritten_entry = dict(file_entry)
                raw_path = str(rewritten_entry.get("file_path") or "").strip()
                if raw_path:
                    rewritten_path = transform(raw_path)
                    rewritten_entry["file_path"] = rewritten_path
                    rewritten_entry["file_name"] = Path(rewritten_path).name
                rewritten_files.append(rewritten_entry)
            updated["files"] = rewritten_files
            rewritten_excluded: list[str] = []
            for raw_path in updated.get("excluded_files", []):
                path_text = str(raw_path).strip()
                if path_text:
                    rewritten_excluded.append(transform(path_text))
            if "excluded_files" in updated or rewritten_excluded:
                updated["excluded_files"] = rewritten_excluded
            return updated

        if kind == "circuit":
            rewritten_instances: list[dict] = []
            for instance in updated.get("instances", []):
                if not isinstance(instance, dict):
                    continue
                rewritten_instance = dict(instance)
                if str(rewritten_instance.get("block_kind", "touchstone")) == "touchstone":
                    raw_path = str(rewritten_instance.get("source_file_id") or "").strip()
                    if raw_path:
                        rewritten_instance["source_file_id"] = transform(raw_path)
                rewritten_instances.append(rewritten_instance)
            updated["instances"] = rewritten_instances

        return updated

    def _resolve_project_payload_touchstone_paths(self, payload: dict, project_dir: Path) -> dict:
        resolved = deepcopy(payload)

        loaded_files: list[dict] = []
        for item in resolved.get("loaded_files", []):
            if not isinstance(item, dict):
                continue
            rewritten_item = dict(item)
            raw_path = str(rewritten_item.get("file_path") or "").strip()
            if raw_path:
                resolved_path = self._resolve_project_reference_path(project_dir, raw_path)
                rewritten_item["file_path"] = resolved_path
                rewritten_item["file_name"] = Path(resolved_path).name
            loaded_files.append(rewritten_item)
        resolved["loaded_files"] = loaded_files

        registry = resolved.get("window_registry")
        if isinstance(registry, dict):
            for kind in ("sp", "tdr", "circuit"):
                raw_entries = registry.get(kind, [])
                if not isinstance(raw_entries, list):
                    continue
                rewritten_entries: list[dict] = []
                for entry in raw_entries:
                    if not isinstance(entry, dict):
                        continue
                    rewritten_entry = dict(entry)
                    state = rewritten_entry.get("state")
                    if isinstance(state, dict):
                        rewritten_entry["state"] = self._rewrite_window_state_touchstone_paths(
                            kind,
                            state,
                            lambda raw, project_dir=project_dir: self._resolve_project_reference_path(project_dir, raw),
                        )
                    rewritten_entries.append(rewritten_entry)
                registry[kind] = rewritten_entries

        for key, kind in (("plots", "sp"), ("tdr_plots", "tdr"), ("circuits", "circuit")):
            raw_states = resolved.get(key, [])
            if not isinstance(raw_states, list):
                continue
            rewritten_states: list[dict] = []
            for state in raw_states:
                if not isinstance(state, dict):
                    continue
                rewritten_states.append(
                    self._rewrite_window_state_touchstone_paths(
                        kind,
                        state,
                        lambda raw, project_dir=project_dir: self._resolve_project_reference_path(project_dir, raw),
                    )
                )
            resolved[key] = rewritten_states

        return resolved

    def _collect_referenced_touchstone_paths(self, payload: dict) -> list[str]:
        seen: set[str] = set()
        referenced: list[str] = []

        def _remember(raw_path: str) -> None:
            normalized = self._normalize_path(raw_path)
            if not normalized or normalized in seen:
                return
            seen.add(normalized)
            referenced.append(normalized)

        registry = payload.get("window_registry")
        if not isinstance(registry, dict):
            return referenced

        for kind in ("sp", "tdr"):
            raw_entries = registry.get(kind, [])
            if not isinstance(raw_entries, list):
                continue
            for entry in raw_entries:
                if not isinstance(entry, dict):
                    continue
                state = entry.get("state")
                if not isinstance(state, dict):
                    continue
                for file_entry in state.get("files", []):
                    if not isinstance(file_entry, dict):
                        continue
                    raw_path = str(file_entry.get("file_path") or "").strip()
                    if raw_path:
                        _remember(raw_path)

        raw_circuits = registry.get("circuit", [])
        if isinstance(raw_circuits, list):
            for entry in raw_circuits:
                if not isinstance(entry, dict):
                    continue
                state = entry.get("state")
                if not isinstance(state, dict):
                    continue
                for instance in state.get("instances", []):
                    if not isinstance(instance, dict):
                        continue
                    if str(instance.get("block_kind", "touchstone")) != "touchstone":
                        continue
                    raw_path = str(instance.get("source_file_id") or "").strip()
                    if raw_path:
                        _remember(raw_path)

        return referenced

    def _copy_touchstone_sources_for_export(
        self,
        source_paths: list[str],
        source_dir: Path,
    ) -> dict[str, str]:
        source_dir.mkdir(parents=True, exist_ok=True)
        copied: dict[str, str] = {}
        used_names: set[str] = set()

        for raw_path in source_paths:
            normalized = self._normalize_path(raw_path)
            if not normalized:
                continue
            source_path = Path(normalized)
            if not source_path.exists():
                raise FileNotFoundError(normalized)

            candidate_name = source_path.name
            name_key = candidate_name.lower()
            if name_key in used_names:
                stem = source_path.stem
                suffix = source_path.suffix
                index = 2
                while True:
                    candidate_name = f"{stem}_{index}{suffix}"
                    name_key = candidate_name.lower()
                    if name_key not in used_names:
                        break
                    index += 1
            used_names.add(name_key)

            target_path = source_dir / candidate_name
            shutil.copy2(source_path, target_path)
            copied[normalized] = (Path(source_dir.name) / candidate_name).as_posix()

        return copied

    def _build_export_project_payload(
        self,
        payload: dict,
        referenced_paths: list[str],
        source_path_map: dict[str, str],
    ) -> dict:
        exported = deepcopy(payload)
        exported["loaded_files"] = [
            {
                "file_path": source_path_map[normalized],
                "file_name": Path(source_path_map[normalized]).name,
            }
            for normalized in referenced_paths
            if normalized in source_path_map
        ]

        registry = exported.get("window_registry")
        if isinstance(registry, dict):
            for kind in ("sp", "tdr", "circuit"):
                raw_entries = registry.get(kind, [])
                if not isinstance(raw_entries, list):
                    continue
                rewritten_entries: list[dict] = []
                for entry in raw_entries:
                    if not isinstance(entry, dict):
                        continue
                    rewritten_entry = dict(entry)
                    state = rewritten_entry.get("state")
                    if isinstance(state, dict):
                        rewritten_entry["state"] = self._rewrite_window_state_touchstone_paths(
                            kind,
                            state,
                            lambda raw, source_path_map=source_path_map: source_path_map.get(self._normalize_path(raw), raw),
                        )
                    rewritten_entries.append(rewritten_entry)
                registry[kind] = rewritten_entries

        for key, kind in (("plots", "sp"), ("tdr_plots", "tdr"), ("circuits", "circuit")):
            raw_states = exported.get(key, [])
            if not isinstance(raw_states, list):
                continue
            rewritten_states: list[dict] = []
            for state in raw_states:
                if not isinstance(state, dict):
                    continue
                rewritten_states.append(
                    self._rewrite_window_state_touchstone_paths(
                        kind,
                        state,
                        lambda raw, source_path_map=source_path_map: source_path_map.get(self._normalize_path(raw), raw),
                    )
                )
            exported[key] = rewritten_states

        return exported

    def _build_export_project_bundle(self, project_path: Path) -> tuple[Path, Path, Path, list[str]]:
        self._project_data_dir = project_path.with_name(f"{project_path.stem}_Data")
        export_warnings = self._export_project_data_files(self._project_data_dir)
        payload = self._build_project_payload()
        referenced_paths = self._collect_referenced_touchstone_paths(payload)

        export_dir = project_path.with_name(f"{project_path.stem}_Export")
        export_project_path = export_dir / project_path.name
        export_data_dir = export_dir / f"{project_path.stem}_Data"
        export_source_dir = export_dir / f"{project_path.stem}_Source"
        export_zip_path = export_dir.with_suffix(".zip")

        try:
            if export_dir.exists():
                shutil.rmtree(export_dir)
            export_dir.mkdir(parents=True, exist_ok=True)

            if export_zip_path.exists():
                export_zip_path.unlink()

            if self._project_data_dir.exists():
                shutil.copytree(self._project_data_dir, export_data_dir)
            else:
                export_data_dir.mkdir(parents=True, exist_ok=True)

            source_path_map = self._copy_touchstone_sources_for_export(referenced_paths, export_source_dir)
            export_payload = self._build_export_project_payload(payload, referenced_paths, source_path_map)
            with open(export_project_path, "w", encoding="utf-8") as fp:
                json.dump(export_payload, fp, indent=2)

            archive_path = shutil.make_archive(
                str(export_dir),
                "zip",
                root_dir=export_dir.parent,
                base_dir=export_dir.name,
            )
        except Exception:
            if export_dir.exists():
                shutil.rmtree(export_dir, ignore_errors=True)
            if export_zip_path.exists():
                try:
                    export_zip_path.unlink()
                except OSError:
                    pass
            raise

        return export_dir, export_project_path, Path(archive_path), export_warnings

    def _export_project(self) -> bool:
        if self._project_path is None:
            QMessageBox.information(
                self,
                "Export Project",
                "Save the project first. A Save Project As dialog will open now.",
            )
            if not self._save_project_as():
                return False

        project_path = Path(str(self._project_path))
        try:
            export_dir, export_project_path, export_zip_path, export_warnings = self._build_export_project_bundle(project_path)
        except FileNotFoundError as exc:
            QMessageBox.critical(
                self,
                "Export failed",
                f"Could not export the project because this Touchstone file is missing:\n{exc}",
            )
            return False
        except OSError as exc:
            QMessageBox.critical(self, "Export failed", f"Could not export project:\n{exc}")
            return False

        message = (
            f"Export folder created:\n{export_dir}\n\n"
            f"Exported project file:\n{export_project_path}\n\n"
            f"ZIP archive:\n{export_zip_path}"
        )
        if export_warnings:
            message += "\n\nSome project data exports were missing:\n" + "\n".join(export_warnings)
        QMessageBox.information(self, "Project exported", message)
        return True

    def _save_project_to_path(self, file_path: str) -> bool:
        self._project_data_dir = Path(file_path).with_name(Path(file_path).stem + "_Data")
        export_errors = self._export_project_data_files(self._project_data_dir)
        payload = self._build_project_payload()

        try:
            with open(file_path, "w", encoding="utf-8") as fp:
                json.dump(payload, fp, indent=2)
        except OSError as exc:
            QMessageBox.critical(self, "Save failed", f"Could not save project:\n{exc}")
            return False

        self._project_path = file_path
        self._project_dirty = False
        self._push_recent_project(file_path)
        self._refresh_project_tree()
        if export_errors:
            QMessageBox.information(
                self,
                "Project saved",
                f"Project saved to:\n{file_path}\n\n"
                f"Data folder: {self._project_data_dir}\n\n"
                "Some exports failed:\n" + "\n".join(export_errors),
            )
        else:
            QMessageBox.information(
                self,
                "Project saved",
                f"Project saved to:\n{file_path}\n\nData folder: {self._project_data_dir}",
            )
        return True

    def _sanitize_file_stem(self, text: str, fallback: str) -> str:
        safe = "".join(ch if (ch.isalnum() or ch in {"-", "_"}) else "_" for ch in text).strip("_")
        return safe or fallback

    def _render_sparameter_plot_pixmap_from_text(self, touchstone_text: str, title: str) -> QPixmap | None:
        try:
            parsed = parse_touchstone_string(touchstone_text, source_name=title)
        except Exception:
            return None
        freqs = np.array(parsed.magnitude_table.frequencies_hz, dtype=float)
        if freqs.size < 2:
            return None
        trace_names = list(parsed.trace_names)
        if not trace_names:
            return None

        selected = trace_names[: min(6, len(trace_names))]
        traces: list[np.ndarray] = []
        for name in selected:
            values = np.array(parsed.magnitude_table.traces_db.get(name, []), dtype=float)
            if values.size != freqs.size:
                continue
            traces.append(values)
        if not traces:
            return None

        width, height = 1200, 760
        left, right, top, bottom = 90, 30, 50, 70
        plot_w = width - left - right
        plot_h = height - top - bottom

        pixmap = QPixmap(width, height)
        pixmap.fill(QColor("#ffffff"))
        painter = QPainter(pixmap)
        painter.setRenderHint(QPainter.Antialiasing, True)

        x0, y0 = left, top
        x1, y1 = left + plot_w, top + plot_h
        painter.setPen(QPen(QColor("#d1d5db"), 1))
        for i in range(6):
            y = int(y0 + i * (plot_h / 5.0))
            painter.drawLine(x0, y, x1, y)
        for i in range(6):
            x = int(x0 + i * (plot_w / 5.0))
            painter.drawLine(x, y0, x, y1)

        all_vals = np.concatenate(traces)
        finite = np.isfinite(all_vals)
        if not np.any(finite):
            painter.end()
            return pixmap
        y_min = float(np.min(all_vals[finite]))
        y_max = float(np.max(all_vals[finite]))
        if y_max <= y_min:
            y_max = y_min + 1.0
        x_min = float(np.min(freqs))
        x_max = float(np.max(freqs))
        if x_max <= x_min:
            x_max = x_min + 1.0

        colors = [
            QColor("#1f77b4"), QColor("#d62728"), QColor("#2ca02c"),
            QColor("#ff7f0e"), QColor("#17becf"), QColor("#9467bd"),
        ]

        def _map_x(v: float) -> float:
            return x0 + ((v - x_min) / (x_max - x_min)) * plot_w

        def _map_y(v: float) -> float:
            return y1 - ((v - y_min) / (y_max - y_min)) * plot_h

        for idx, values in enumerate(traces):
            painter.setPen(QPen(colors[idx % len(colors)], 2))
            prev_x = _map_x(float(freqs[0]))
            prev_y = _map_y(float(values[0]))
            for i in range(1, freqs.size):
                x = _map_x(float(freqs[i]))
                y = _map_y(float(values[i]))
                painter.drawLine(int(prev_x), int(prev_y), int(x), int(y))
                prev_x, prev_y = x, y

        painter.setPen(QPen(QColor("#111827"), 2))
        painter.drawRect(x0, y0, plot_w, plot_h)
        painter.drawText(20, 26, title)
        painter.drawText(20, height - 20, "Frequency [Hz]")
        painter.drawText(width - 180, 26, "Magnitude [dB]")
        painter.end()
        return pixmap

    def _export_project_data_files(self, data_dir: Path) -> list[str]:
        errors: list[str] = []
        try:
            data_dir.mkdir(parents=True, exist_ok=True)
        except OSError as exc:
            return [f"Could not create data folder '{data_dir}': {exc}"]

        for entry_id, entry in self._window_registry["circuit"].items():
            circuit_name = str(entry.get("title", f"Circuit {entry_id}"))
            sim_mode = "S-Parameters"
            state_dict = entry.get("state")
            if isinstance(state_dict, dict):
                sim_mode = str(state_dict.get("simulation_mode", "S-Parameters"))

            if sim_mode == "S-Parameters":
                entry["output_kind"] = "sparam"
                entry["eye_file"] = None
                sparam_file = str(entry.get("sparam_file") or "").strip()
                if sparam_file and not (data_dir / sparam_file).exists():
                    errors.append(f"{circuit_name}: missing referenced SnP file ({sparam_file})")
                    entry["sparam_file"] = None
                sparam_plot_file = str(entry.get("sparam_plot_file") or "").strip()
                if sparam_plot_file and not (data_dir / sparam_plot_file).exists():
                    errors.append(f"{circuit_name}: missing referenced S-plot file ({sparam_plot_file})")
                    entry["sparam_plot_file"] = None
            else:
                entry["output_kind"] = "eye"
                entry["sparam_file"] = None
                entry["sparam_plot_file"] = None
                eye_file = str(entry.get("eye_file") or "").strip()
                if eye_file and not (data_dir / eye_file).exists():
                    errors.append(f"{circuit_name}: missing referenced Eye file ({eye_file})")
                    entry["eye_file"] = None

        self._refresh_project_tree()
        return errors

    def _save_project_as(self) -> bool:
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Project As",
            self._project_path or "SPUtility_project.json",
            "SPUtility Project (*.json);;All files (*)",
        )
        if not file_path:
            return False

        return self._save_project_to_path(file_path)

    def _save_project(self) -> bool:
        if self._project_path is None:
            return self._save_project_as()

        return self._save_project_to_path(self._project_path)

    def _confirm_project_close(self) -> bool:
        if not self._project_dirty:
            return True

        box = QMessageBox(self)
        box.setWindowTitle("Unsaved Changes")
        box.setText("Project has unsaved changes. Do you want to save before closing?")
        box.setIcon(QMessageBox.Warning)
        box.setStandardButtons(
            QMessageBox.Save | QMessageBox.Discard | QMessageBox.Cancel
        )
        box.setDefaultButton(QMessageBox.Save)

        choice = box.exec()
        if choice == QMessageBox.Save:
            return self._save_project()

        return choice == QMessageBox.Discard

    def _reset_project_workspace(self) -> None:
        self._windows.close_all()
        self._state.clear_files()
        self._plot_counter = 0
        self._tdr_counter = 0
        self._circuit_counter = 0
        self._project_dirty = False
        self._project_path = None
        self._project_data_dir = None
        self._window_registry = {"sp": {}, "tdr": {}, "circuit": {}}
        self._open_windows = {}
        self._open_window_keys = {}
        self._rebuild_tables_menu()
        self._refresh_project_tree()
        self._on_active_widget_changed(None)

    def _close_project(self) -> None:
        if not self._confirm_project_close():
            return

        self._reset_project_workspace()

    def _open_project(self) -> None:
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Open Project",
            "",
            "SPUtility Project (*.json);;All files (*)",
        )
        if not file_path:
            return

        self._load_project_from_path(file_path)

    def _load_project_from_path(self, file_path: str) -> None:

        try:
            with open(file_path, "r", encoding="utf-8") as fp:
                payload = json.load(fp)
        except (OSError, json.JSONDecodeError) as exc:
            QMessageBox.critical(self, "Open failed", f"Could not open project:\n{exc}")
            return

        payload = self._resolve_project_payload_touchstone_paths(payload, Path(file_path).parent)

        raw_files = payload.get("loaded_files", [])
        file_paths: list[str] = []
        for item in raw_files:
            path = item.get("file_path") if isinstance(item, dict) else None
            if isinstance(path, str) and path:
                file_paths.append(path)

        if not self._confirm_project_close():
            return

        self._reset_project_workspace()

        _, errors = self._state.load_files(file_paths)

        restored_plot = 0
        restored_tdr = 0
        restored_circuit = 0

        registry = payload.get("window_registry")
        if isinstance(registry, dict):
            for kind in ("sp", "tdr", "circuit"):
                raw_entries = registry.get(kind, [])
                if not isinstance(raw_entries, list):
                    continue
                for item in raw_entries:
                    if not isinstance(item, dict):
                        continue
                    entry_id = str(item.get("id") or self._new_entry_id(kind))
                    self._window_registry[kind][entry_id] = {
                        "id": entry_id,
                        "kind": kind,
                        "title": str(item.get("title", "Window")),
                        "state": dict(item.get("state", {})) if isinstance(item.get("state"), dict) else {},
                        "window_number": int(item.get("window_number", 0)),
                        "is_open": bool(item.get("is_open", False)),
                        "window_size": item.get("window_size"),
                        "eye_file": item.get("eye_file"),
                        "sparam_file": item.get("sparam_file"),
                        "sparam_plot_file": item.get("sparam_plot_file"),
                        "transient_file": item.get("transient_file"),
                        "output_kind": item.get("output_kind"),
                        "parent_circuit_entry_id": item.get("parent_circuit_entry_id"),
                    }

            for kind in ("sp", "tdr", "circuit"):
                for entry_id, entry in self._window_registry[kind].items():
                    if bool(entry.get("is_open", False)):
                        self._activate_or_open_window_entry(kind, entry_id)

            restored_plot = sum(1 for e in self._window_registry["sp"].values() if bool(e.get("is_open", False)))
            restored_tdr = sum(1 for e in self._window_registry["tdr"].values() if bool(e.get("is_open", False)))
            restored_circuit = sum(1 for e in self._window_registry["circuit"].values() if bool(e.get("is_open", False)))
        else:
            # Backward compatibility with older project files.
            for plot_state in payload.get("plots", []):
                if not isinstance(plot_state, dict):
                    continue
                self._plot_counter += 1
                plot_win = PlotWindow(self._state, window_number=self._plot_counter)
                plot_win.apply_project_state(plot_state)
                entry_id = self._register_window_entry("sp", plot_win)
                window_size = plot_state.get("window_size")
                if (
                    isinstance(window_size, list)
                    and len(window_size) == 2
                    and all(isinstance(v, int) for v in window_size)
                ):
                    plot_win.resize(window_size[0], window_size[1])
                else:
                    plot_win.resize(*default_child_window_size("plots"))
                plot_win.project_modified.connect(self._mark_project_dirty)
                self._bind_open_window("sp", entry_id, plot_win)
                self._windows.present(plot_win)
                restored_plot += 1

            for tdr_state in payload.get("tdr_plots", []):
                if not isinstance(tdr_state, dict):
                    continue
                self._tdr_counter += 1
                tdr_win = TdrWindow(self._state, window_number=self._tdr_counter)
                tdr_win.apply_project_state(tdr_state)
                entry_id = self._register_window_entry("tdr", tdr_win)
                window_size = tdr_state.get("window_size")
                if (
                    isinstance(window_size, list)
                    and len(window_size) == 2
                    and all(isinstance(v, int) for v in window_size)
                ):
                    tdr_win.resize(window_size[0], window_size[1])
                else:
                    tdr_win.resize(*default_child_window_size("tdr"))
                tdr_win.project_modified.connect(self._mark_project_dirty)
                self._bind_open_window("tdr", entry_id, tdr_win)
                self._windows.present(tdr_win)
                restored_tdr += 1

            for circuit_state in payload.get("circuits", []):
                if not isinstance(circuit_state, dict):
                    continue
                self._circuit_counter += 1
                circuit_win = CircuitWindow(self._state, window_number=self._circuit_counter)
                circuit_win.apply_project_state(circuit_state)
                entry_id = self._register_window_entry("circuit", circuit_win)
                window_size = circuit_state.get("window_size")
                if (
                    isinstance(window_size, list)
                    and len(window_size) == 2
                    and all(isinstance(v, int) for v in window_size)
                ):
                    circuit_win.resize(window_size[0], window_size[1])
                else:
                    circuit_win.resize(*default_child_window_size("circuits"))
                circuit_win.project_modified.connect(self._mark_project_dirty)
                self._bind_open_window("circuit", entry_id, circuit_win)
                self._windows.present(circuit_win)
                restored_circuit += 1

        if errors:
            QMessageBox.warning(
                self,
                "Project loaded with warnings",
                "Some files could not be loaded:\n" + "\n".join(errors),
            )
        else:
            QMessageBox.information(
                self,
                "Project loaded",
                f"Loaded {len(file_paths)} file(s), restored {restored_plot} plot window(s), restored {restored_tdr} TDR window(s), and restored {restored_circuit} circuit window(s).",
            )
        self._project_path = file_path
        self._project_data_dir = Path(file_path).with_name(Path(file_path).stem + "_Data")
        self._project_dirty = False
        self._push_recent_project(file_path)
        self._refresh_project_tree()

    # ── Tables menu ───────────────────────────────────────────────────────

    def _rebuild_tables_menu(self) -> None:
        self._tables_menu.clear()
        loaded_files = self._state.get_loaded_files()
        if not loaded_files:
            a = self._tables_menu.addAction("No files loaded")
            a.setEnabled(False)
            return

        touchstone_icon = self._icon_for_kind("touchstone-file")
        for loaded in loaded_files:
            submenu = self._tables_menu.addMenu(loaded.display_name)
            submenu.setIcon(touchstone_icon)
            submenu.addAction(
                "Raw data table",
                lambda checked=False, item=loaded: self._show_raw_table(item),
            )
            submenu.addAction(
                "Magnitude table [dB]",
                lambda checked=False, item=loaded: self._show_magnitude_table(item),
            )
            submenu.addSeparator()
            submenu.addAction(
                "Unload file",
                lambda checked=False, file_id=loaded.file_id: self._unload_file(file_id),
            )

    def _unload_file(self, file_id: str) -> None:
        loaded = self._state.get_file(file_id)
        if loaded is None:
            return

        for widget in self._windows.list_widgets():
            if isinstance(widget, CircuitWindow) and widget.references_file(file_id):
                QMessageBox.warning(
                    self,
                    "File in use",
                    "This file is used by at least one circuit window. Remove its block instances before unloading it.",
                )
                return

        for widget in list(self._windows.list_widgets()):
            if isinstance(widget, TableWindow) and widget.file_id == file_id:
                widget.close()

        removed = self._state.unload_file(file_id)
        if removed:
            self._mark_project_dirty()
            self._refresh_project_tree()

    def _show_raw_table(self, loaded: LoadedTouchstone) -> None:
        win = TableWindow(
            f"{loaded.display_name} - Raw data",
            RawDataTableModel(loaded.data),
            file_id=loaded.file_id,
        )
        win.resize(*default_child_window_size("tables"))
        self._windows.present(win)

    def _show_magnitude_table(self, loaded: LoadedTouchstone) -> None:
        win = TableWindow(
            f"{loaded.display_name} - Magnitude [dB]",
            MagnitudeTableModel(loaded.data),
            file_id=loaded.file_id,
        )
        win.resize(*default_child_window_size("tables"))
        self._windows.present(win)

    # ── Plot window ───────────────────────────────────────────────────────

    def _open_plot_window(self) -> None:
        self._plot_counter += 1
        plot_win = PlotWindow(self._state, window_number=self._plot_counter)
        plot_win.project_modified.connect(self._mark_project_dirty)
        plot_win.resize(*default_child_window_size("plots"))
        entry_id = self._register_window_entry("sp", plot_win)
        self._bind_open_window("sp", entry_id, plot_win)
        self._windows.present(plot_win)
        self._mark_project_dirty()
        self._refresh_project_tree()

    def _open_tdr_window(self) -> None:
        self._tdr_counter += 1
        tdr_win = TdrWindow(self._state, window_number=self._tdr_counter)
        tdr_win.project_modified.connect(self._mark_project_dirty)
        tdr_win.resize(*default_child_window_size("tdr"))
        entry_id = self._register_window_entry("tdr", tdr_win)
        self._bind_open_window("tdr", entry_id, tdr_win)
        self._windows.present(tdr_win)
        self._mark_project_dirty()
        self._refresh_project_tree()

    def _open_circuit_window(self) -> None:
        circuit_number = self._next_available_circuit_number()
        self._circuit_counter = max(self._circuit_counter, circuit_number)
        circuit_win = CircuitWindow(self._state, window_number=circuit_number)
        circuit_win.project_modified.connect(self._mark_project_dirty)
        circuit_win.resize(*default_child_window_size("circuits"))
        entry_id = self._register_window_entry("circuit", circuit_win)
        self._bind_open_window("circuit", entry_id, circuit_win)
        self._windows.present(circuit_win)
        self._mark_project_dirty()
        self._refresh_project_tree()

    def _next_available_circuit_number(self) -> int:
        used_numbers: set[int] = set()
        for widget in self._windows.list_widgets():
            if isinstance(widget, CircuitWindow):
                used_numbers.add(int(getattr(widget, "window_number", 0)))
        for entry in self._window_registry["circuit"].values():
            used_numbers.add(int(entry.get("window_number", 0)))
        number = 1
        while number in used_numbers:
            number += 1
        return number

    def _duplicate_active_circuit(self) -> None:
        active_widget = self._windows.active_widget()
        if not isinstance(active_widget, CircuitWindow):
            QMessageBox.information(
                self,
                "Duplicate circuit",
                "Select a Circuit Composer window before using this command.",
            )
            return

        source_window = active_widget
        source_state = source_window.export_project_state()
        circuit_number = self._next_available_circuit_number()
        self._circuit_counter = max(self._circuit_counter, circuit_number)

        duplicated_window = CircuitWindow(self._state, window_number=circuit_number)
        duplicated_window.apply_project_state(source_state)
        duplicated_window.project_modified.connect(self._mark_project_dirty)
        duplicated_window.resize(source_window.size())
        entry_id = self._register_window_entry("circuit", duplicated_window)
        self._bind_open_window("circuit", entry_id, duplicated_window)
        self._windows.present(duplicated_window)
        self._mark_project_dirty()
        self._refresh_project_tree()

    def _resize_all_graph_windows(self) -> None:
        graph_windows: list[object] = []
        plot_windows: list[PlotWindow] = []
        for widget in self._windows.list_widgets():
            if isinstance(widget, (PlotWindow, TdrWindow, EyeDiagramWindow)):
                graph_windows.append(widget)
            if isinstance(widget, PlotWindow):
                plot_windows.append(widget)

        if not graph_windows:
            QMessageBox.information(
                self,
                "No graph windows",
                "There are no graph windows to resize.",
            )
            return

        active_widget = self._windows.active_widget()
        if isinstance(active_widget, (PlotWindow, TdrWindow, EyeDiagramWindow)):
            source = active_widget
        elif plot_windows:
            source = max(plot_windows, key=lambda w: w.window_number)
        else:
            source = graph_windows[-1]

        source_size = source.size()
        state = source.get_graph_layout_state() if isinstance(source, PlotWindow) else None

        for widget in graph_windows:
            if widget is source:
                continue

            widget.resize(source_size)
            if isinstance(source, PlotWindow) and isinstance(widget, PlotWindow) and isinstance(state, dict):
                widget.apply_graph_layout_state(state)

    # ── Report export helpers ────────────────────────────────────────────

    def _ensure_docx_support(self):
        try:
            from docx import Document
            from docx.shared import Inches
        except Exception as exc:
            QMessageBox.warning(
                self,
                "Report unavailable",
                "python-docx is required to create Word reports. "
                f"Install dependency and retry.\nDetails: {exc}",
            )
            return None, None
        return Document, Inches

    def _choose_report_path(self, title: str, default_name: str) -> str | None:
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            title,
            default_name,
            "Word Document (*.docx);;All files (*)",
        )
        if not file_path:
            return None
        if not file_path.lower().endswith(".docx"):
            file_path = f"{file_path}.docx"
        return file_path

    def _pixmap_to_png_stream(self, pixmap) -> BytesIO | None:
        if pixmap is None or pixmap.isNull():
            return None
        raw = QByteArray()
        buf = QBuffer(raw)
        if not buf.open(QIODevice.WriteOnly):
            return None
        ok = pixmap.save(buf, "PNG")
        buf.close()
        if not ok:
            return None
        stream = BytesIO(bytes(raw))
        stream.seek(0)
        return stream

    def _report_header(self, doc, title: str) -> None:
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    def _create_sp_plots_report(self) -> None:
        windows = list(self._windows.widgets_of_type(PlotWindow))
        if not windows:
            QMessageBox.information(self, "Report", "No SP Plot windows are currently open.")
            return

        Document, Inches = self._ensure_docx_support()
        if Document is None:
            return
        output_path = self._choose_report_path(
            "Create SP Plots Reports",
            "SP_Plots_Report.docx",
        )
        if output_path is None:
            return

        doc = Document()
        self._report_header(doc, "SP Plots Report")

        for idx, win in enumerate(windows, start=1):
            if idx > 1:
                doc.add_page_break()
            doc.add_heading(f"Plot Window {idx}: {win.windowTitle()}", level=2)
            state = win.export_project_state()
            plotted_files: list[str] = []
            for file_entry in state.get("files", []):
                selected = file_entry.get("selected_parameters", [])
                if not isinstance(selected, list) or len(selected) == 0:
                    continue
                file_path = str(file_entry.get("file_path", "")).strip()
                if not file_path:
                    continue
                plotted_files.append(Path(file_path).name)
            plotted_files = list(dict.fromkeys(plotted_files))
            doc.add_paragraph(
                "Files: " + (", ".join(plotted_files) if plotted_files else "n/a")
            )

            stream = self._pixmap_to_png_stream(win._plot_widget.grab())
            if stream is not None:
                doc.add_picture(stream, width=Inches(6.8))

        try:
            doc.save(output_path)
        except Exception as exc:
            QMessageBox.warning(self, "Report", f"Could not save report:\n{exc}")
            return
        QMessageBox.information(self, "Report", f"SP plots report created:\n{output_path}")

    def _create_tdr_plots_report(self) -> None:
        windows = list(self._windows.widgets_of_type(TdrWindow))
        if not windows:
            QMessageBox.information(self, "Report", "No TDR Plot windows are currently open.")
            return

        Document, Inches = self._ensure_docx_support()
        if Document is None:
            return
        output_path = self._choose_report_path(
            "Create TDR plots reports",
            "TDR_Plots_Report.docx",
        )
        if output_path is None:
            return

        doc = Document()
        self._report_header(doc, "TDR Plots Report")

        for idx, win in enumerate(windows, start=1):
            if idx > 1:
                doc.add_page_break()
            doc.add_heading(f"TDR Window {idx}: {win.windowTitle()}", level=2)
            state = win.export_project_state()
            plotted_files: list[str] = []
            for file_entry in state.get("files", []):
                if not bool(file_entry.get("show", True)):
                    continue
                selected_trace = str(file_entry.get("selected_trace", "")).strip()
                if not selected_trace:
                    continue
                file_path = str(file_entry.get("file_path", "")).strip()
                if not file_path:
                    continue
                plotted_files.append(Path(file_path).name)
            plotted_files = list(dict.fromkeys(plotted_files))
            doc.add_paragraph(
                "Files: " + (", ".join(plotted_files) if plotted_files else "n/a")
            )

            stream = self._pixmap_to_png_stream(win._plot_widget.grab())
            if stream is not None:
                doc.add_picture(stream, width=Inches(6.8))

        try:
            doc.save(output_path)
        except Exception as exc:
            QMessageBox.warning(self, "Report", f"Could not save report:\n{exc}")
            return
        QMessageBox.information(self, "Report", f"TDR plots report created:\n{output_path}")

    def _create_eye_diagrams_report(self) -> None:
        eye_entries: list[tuple[CircuitWindow, object]] = []
        generation_issues: list[str] = []
        circuit_windows = list(self._windows.widgets_of_type(CircuitWindow))

        if not circuit_windows:
            QMessageBox.information(self, "Report", "No Circuit windows are currently open.")
            return

        for widget in circuit_windows:
            windows = widget.get_open_eye_windows()
            if not windows:
                generated, err = widget.generate_eye_window_for_report()
                if generated is not None:
                    windows = [generated]
                elif err is not None:
                    generation_issues.append(f"{widget.circuit_display_name()}: {err}")
            for eye_win in windows:
                eye_entries.append((widget, eye_win))

        if not eye_entries:
            details = "\n".join(generation_issues)
            msg = "No Eye diagrams could be generated for the current project windows."
            if details:
                msg += "\n\n" + details
            QMessageBox.information(self, "Report", msg)
            return

        Document, Inches = self._ensure_docx_support()
        if Document is None:
            return
        output_path = self._choose_report_path(
            "Create EYE diagrams report",
            "EYE_Diagrams_Report.docx",
        )
        if output_path is None:
            return

        doc = Document()
        self._report_header(doc, "Eye Diagrams Report")

        def _fmt_mv(v: float) -> str:
            import math as _math
            return "n/a" if not _math.isfinite(v) else f"{v * 1000:.2f} mV"

        def _fmt_ps(v: float) -> str:
            import math as _math
            return "n/a" if not _math.isfinite(v) else f"{v:.2f} ps"

        def _fmt_pct(v: float) -> str:
            import math as _math
            return "n/a" if not _math.isfinite(v) else f"{v:.2f} %"

        for idx, (circuit_win, eye_win) in enumerate(eye_entries, start=1):
            if idx > 1:
                doc.add_page_break()

            state = eye_win.export_report_state()
            meas = state.get("measurements", {})
            files = circuit_win.report_touchstone_file_names()
            circuit_name = circuit_win.circuit_display_name()

            doc.add_heading(f"Eye Diagram {idx}: {eye_win.windowTitle()}", level=2)
            doc.add_paragraph("Circuit: " + circuit_name)
            doc.add_paragraph("Files: " + (", ".join(files) if files else "n/a"))

            settings_lines = [
                f"Mode: {state.get('mode', 'n/a')}",
                f"Bitrate: {float(state.get('bitrate_gbps', 0.0)):.3f} Gbps",
                f"Encoding: {state.get('encoding', 'n/a')}",
                f"PRBS: {state.get('prbs_pattern', 'n/a')}",
                f"Bits: {int(state.get('num_bits', 0))}",
                f"Rise/Fall: {float(state.get('rise_time_ps', 0.0)):.2f} ps / {float(state.get('fall_time_ps', 0.0)):.2f} ps",
                f"Vhigh/Vlow: {float(state.get('voltage_high_v', 0.0)):.4f} V / {float(state.get('voltage_low_v', 0.0)):.4f} V",
                f"Eye span: {int(state.get('eye_span_ui', 0))} UI",
                f"Render mode: {state.get('render_mode', 'n/a')}",
                f"Quality preset: {state.get('quality_preset', 'n/a')}",
                f"Statistical: {'ON' if bool(state.get('statistical_enabled', False)) else 'OFF'}",
                f"Noise RMS: {float(state.get('noise_rms_mv', 0.0)):.2f} mV",
                f"Jitter RMS: {float(state.get('jitter_rms_ps', 0.0)):.2f} ps",
            ]
            doc.add_paragraph("Settings:\n" + "\n".join(settings_lines))

            if isinstance(meas, dict):
                m_lines = [
                    f"One Level: {_fmt_mv(float(meas.get('one_level', float('nan'))))}",
                    f"Zero Level: {_fmt_mv(float(meas.get('zero_level', float('nan'))))}",
                    f"Eye Amplitude: {_fmt_mv(float(meas.get('eye_amplitude', float('nan'))))}",
                    f"Eye Height: {_fmt_mv(float(meas.get('eye_height', float('nan'))))}",
                    f"Eye Width: {_fmt_ps(float(meas.get('eye_width_ps', meas.get('width_ps', float('nan')))))}",
                    f"Eye Crossing: {_fmt_pct(float(meas.get('eye_crossing_pct', float('nan'))))}",
                    f"DJ (pp): {_fmt_ps(float(meas.get('dj_pp_s', float('nan'))) * 1e12)}",
                    f"RJ (σ): {_fmt_ps(float(meas.get('sigma_rj_s', float('nan'))) * 1e12)}",
                    f"TJ (pp@BER): {_fmt_ps(float(meas.get('tj_pp_s', float('nan'))) * 1e12)}",
                    f"Bit Period: {_fmt_ps(float(meas.get('bit_period_ps', float('nan'))))}",
                ]
            else:
                m_lines = ["Measurements: n/a"]
            doc.add_paragraph("Measurements:\n" + "\n".join(m_lines))

            eye_stream = self._pixmap_to_png_stream(eye_win.grab_eye_plot_pixmap())
            if eye_stream is not None:
                doc.add_paragraph("Eye Diagram Snapshot")
                doc.add_picture(eye_stream, width=Inches(6.8))

            circuit_stream = self._pixmap_to_png_stream(circuit_win.grab_circuit_snapshot_pixmap())
            if circuit_stream is not None:
                doc.add_paragraph("Related Circuit Snapshot")
                doc.add_picture(circuit_stream, width=Inches(6.8))

        try:
            doc.save(output_path)
        except Exception as exc:
            QMessageBox.warning(self, "Report", f"Could not save report:\n{exc}")
            return

        if generation_issues:
            QMessageBox.information(
                self,
                "Report",
                f"Eye diagrams report created:\n{output_path}\n\n"
                "Some circuit windows could not generate an eye diagram:\n"
                + "\n".join(generation_issues),
            )
        else:
            QMessageBox.information(self, "Report", f"Eye diagrams report created:\n{output_path}")

    # ── View helpers ──────────────────────────────────────────────────────

    def _minimize_all(self) -> None:
        self._windows.minimize_all()

    def _restore_all(self) -> None:
        self._windows.restore_all()

    # ── Help > About ──────────────────────────────────────────────────────

    def _show_help(self) -> None:
        try:
            help_resource = resources.files("sparams_utility.resources.help").joinpath(
                "help_en.html"
            )
        except (FileNotFoundError, ModuleNotFoundError, OSError) as exc:
            QMessageBox.warning(
                self,
                "Help unavailable",
                f"Could not load the user guide:\n{exc}",
            )
            return

        dlg = QDialog(self)
        dlg.setWindowTitle(f"{pkg.__app_name__} User Guide")
        dlg.resize(820, 640)

        layout = QVBoxLayout(dlg)
        browser = QTextBrowser(dlg)
        browser.setOpenExternalLinks(True)
        browser.setReadOnly(True)
        with resources.as_file(help_resource) as help_path:
            browser.setSource(QUrl.fromLocalFile(str(help_path)))
            layout.addWidget(browser)
            dlg.exec()

    def _show_about(self) -> None:
        dlg = QDialog(self)
        dlg.setWindowTitle(f"About {pkg.__app_name__}")
        dlg.setFixedSize(340, 200)
        layout = QVBoxLayout(dlg)
        layout.setSpacing(8)

        for html in (
            f"<b style='font-size:16px'>{pkg.__app_name__}</b>",
            f"Version {pkg.__version__}",
            "<hr>",
            f"<b>Author:</b> {pkg.__author__}",
            "",
            "S-Parameters Touchstone Viewer",
        ):
            lbl = QLabel(html)
            lbl.setAlignment(Qt.AlignCenter)
            lbl.setTextFormat(Qt.RichText)
            layout.addWidget(lbl)

        dlg.exec()

    def _mark_project_dirty(self) -> None:
        self._project_dirty = True

    def closeEvent(self, event: QCloseEvent) -> None:  # noqa: N802
        if self._confirm_project_close():
            self._windows.close_all()
            event.accept()
            return

        event.ignore()

    # ── Child window manager exposure & menu wiring ──────────────────────

    def child_window_manager(self):
        return self._windows

    def _on_active_widget_changed(self, widget) -> None:
        self._active_circuit_menu.menuAction().setVisible(isinstance(widget, CircuitWindow))
        self._refresh_project_tree()

    def _populate_window_menu(self) -> None:
        self._window_menu.clear()
        order = ("plots", "tdr", "eye", "transient", "circuits", "tables", "misc")
        category_titles = {
            "plots": "S-Parameter Plots",
            "tdr": "TDR Plots",
            "eye": "Eye Diagrams",
            "transient": "Transient Results",
            "circuits": "Circuits",
            "tables": "Tables",
            "misc": "Other Windows",
        }
        any_added = False
        for category in order:
            widgets = self._windows.list_widgets(category)
            if not widgets:
                continue
            if any_added:
                self._window_menu.addSeparator()
            header = self._window_menu.addAction(f"— {category_titles[category]} —")
            header.setEnabled(False)
            for widget in widgets:
                title = widget.windowTitle() or "Untitled"
                action = self._window_menu.addAction(title)
                action.triggered.connect(
                    lambda _checked=False, w=widget: self._windows.set_active_widget(w)
                )
            any_added = True
        if not any_added:
            empty = self._window_menu.addAction("(no open windows)")
            empty.setEnabled(False)
        # Detached windows actions
        detached = self._windows.detached_widgets()
        if detached:
            self._window_menu.addSeparator()
            header = self._window_menu.addAction("— Detached —")
            header.setEnabled(False)
            for widget in detached:
                title = widget.windowTitle() or "Untitled"
                act = self._window_menu.addAction(f"{title} (detached)")
                act.triggered.connect(
                    lambda _c=False, w=widget: self._windows.set_active_widget(w)
                )
            reattach = self._window_menu.addAction("Reattach all detached")
            reattach.triggered.connect(self._windows.reattach_all)
